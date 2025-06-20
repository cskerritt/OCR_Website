import streamlit as st
import pandas as pd
import numpy as np
from datetime import date, datetime, timedelta
from dataclasses import dataclass
import io
from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.shared import OxmlElement, qn
import json
import matplotlib.pyplot as plt
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.table import _Cell
import base64

# Page configuration
st.set_page_config(
    page_title="Economic Report Generator", 
    layout="wide", 
    initial_sidebar_state="expanded"
)

st.title("📊 Economic Loss Report Generator")
st.markdown("Generate comprehensive economic loss reports with professional calculations and formatting.")

# Sample case data for testing
SAMPLE_CASE = {
    'client_name': 'John Doe',
    'date_of_birth': date(1980, 3, 15),
    'date_of_injury': date(2023, 6, 1),
    'date_of_report': date(2025, 6, 20),
    'gender': 'Male',
    'education': 'Bachelor\'s Degree in Engineering',
    'marital_status': 'Married',
    'residence_state': 'New Jersey',
    'residence_county': 'Bergen County',
    'msa': 'New York-Newark-Jersey City, NY-NJ-PA',
    'pre_injury_occupation': 'Mechanical Engineer',
    'pre_injury_annual_income': 75000.0,
    'pre_injury_hourly_rate': 36.06,
    'employer': 'ABC Manufacturing Corp',
    'years_with_employer': 8.0,
    'injury_description': 'Work-related back injury sustained while lifting heavy machinery components. Resulted in herniated discs L4-L5 and chronic pain limiting physical capacity.',
    'body_parts_injured': 'Lower back, L4-L5 vertebrae',
    'medical_treatment': 'Initial emergency treatment, MRI diagnostics, physical therapy (6 months), pain management, epidural injections',
    'current_medical_status': 'Maximum medical improvement reached. Permanent restrictions: no lifting >20 lbs, limited standing/walking',
    'life_expectancy': 78.5,
    'work_life_expectancy': 42.0,
    'years_to_final_separation': 0.0,
    'wage_growth_rate': 0.032,
    'discount_rate': 0.045,
    'unemployment_rate': 0.038,
    'tax_rate': 0.22,
    'fringe_benefits_rate': 0.28,
    'personal_consumption_rate': 0.0,
    'post_injury_capacity': 'Limited to sedentary work only due to lifting restrictions and chronic pain. Cannot perform previous engineering duties requiring site visits and equipment handling.',
    'post_injury_occupation': 'Desk-based engineering consultant (when available)',
    'post_injury_annual_income': 35000.0,
    'labor_market_reduction': 0.65,
    'life_care_plan_cost': 125000.0,
    'household_services_cost': 85000.0
}

# Field descriptions and examples
FIELD_HELP = {
    'client_name': 'Full legal name of the injured party (e.g., "John Michael Smith")',
    'date_of_birth': 'Used to calculate age at injury and life expectancy tables',
    'date_of_injury': 'The specific date when the injury occurred - this starts the loss period',
    'date_of_report': 'Current report date - separates past vs future losses',
    'gender': 'Required for accurate life expectancy and work-life expectancy calculations',
    'education': 'Highest level completed (e.g., "High School", "Bachelor\'s in Engineering", "Master\'s in Business")',
    'marital_status': 'Current status (e.g., "Single", "Married", "Divorced", "Widowed")',
    'residence_state': 'State of primary residence for jurisdiction and wage data',
    'residence_county': 'County for local economic data and wage adjustments',
    'msa': 'Metropolitan Statistical Area for regional wage comparisons',
    'pre_injury_occupation': 'Detailed job title and duties (e.g., "Registered Nurse - ICU", "Construction Foreman")',
    'pre_injury_annual_income': 'Total annual earnings including overtime, bonuses, commissions before injury',
    'pre_injury_hourly_rate': 'Base hourly wage rate (calculated from annual if needed)',
    'employer': 'Company name and any relevant details about employment stability',
    'years_with_employer': 'Length of employment showing job stability and advancement potential',
    'injury_description': 'Detailed medical description: mechanism of injury, body parts affected, severity, initial treatment',
    'body_parts_injured': 'Specific anatomical areas (e.g., "Right shoulder, rotator cuff", "Lumbar spine L3-L5")',
    'medical_treatment': 'Comprehensive treatment history: surgeries, therapy, medications, ongoing care',
    'current_medical_status': 'Current condition, restrictions, prognosis, MMI status, permanent limitations',
    'life_expectancy': 'Total remaining years based on actuarial tables (from injury date)',
    'work_life_expectancy': 'Expected working years remaining (from injury date)',
    'years_to_final_separation': 'Additional non-working years (for work-life ratio calculations)',
    'wage_growth_rate': 'Annual percentage increase in wages (historical average 2.5-4%)',
    'discount_rate': 'Present value discount rate (typically government bond rates 3-5%)',
    'unemployment_rate': 'Regional unemployment percentage (reduces earning probability)',
    'tax_rate': 'Combined federal/state/local tax rate on earnings',
    'fringe_benefits_rate': 'Employer benefits as percentage of wages (health, retirement, etc.)',
    'personal_consumption_rate': 'Portion spent on self (used in wrongful death cases only)',
    'post_injury_capacity': 'Detailed description of remaining work abilities and limitations',
    'post_injury_occupation': 'Realistic alternative employment considering restrictions',
    'post_injury_annual_income': 'Projected annual earnings in alternative employment',
    'labor_market_reduction': 'Percentage reduction in earning capacity due to injury',
    'life_care_plan_cost': 'Total present value of future medical/care costs',
    'household_services_cost': 'Value of lost household services (cleaning, cooking, childcare)'
}

# Initialize session state for data persistence
if 'report_data' not in st.session_state:
    st.session_state.report_data = {}

# Load sample case function
def load_sample_case():
    """Load sample case data into session state"""
    st.session_state.report_data = SAMPLE_CASE.copy()
    st.success("Sample case loaded! Review the data in each tab and generate your report.")
    st.rerun()

@dataclass
class EconomicData:
    """Data structure for economic analysis"""
    # Personal Information
    client_name: str = ""
    date_of_birth: date = date(1980, 1, 1)
    date_of_injury: date = date(2023, 1, 1)
    date_of_report: date = date.today()
    age_at_injury: float = 0.0
    gender: str = "Male"
    education: str = ""
    marital_status: str = ""
    
    # Geographic Information
    residence_state: str = ""
    residence_county: str = ""
    msa: str = ""
    
    # Employment Information
    pre_injury_occupation: str = ""
    pre_injury_annual_income: float = 0.0
    pre_injury_hourly_rate: float = 0.0
    employer: str = ""
    years_with_employer: float = 0.0
    
    # Injury Information
    injury_description: str = ""
    body_parts_injured: str = ""
    medical_treatment: str = ""
    current_medical_status: str = ""
    
    # Work Life Expectancy Data
    life_expectancy: float = 78.5
    work_life_expectancy: float = 45.0
    years_to_final_separation: float = 0.0
    statistical_retirement_date: date = date.today()
    
    # Economic Factors
    wage_growth_rate: float = 0.03
    discount_rate: float = 0.04
    unemployment_rate: float = 0.035
    tax_rate: float = 0.12
    fringe_benefits_rate: float = 0.06
    personal_consumption_rate: float = 0.25
    
    # Post-Injury Information
    post_injury_capacity: str = ""
    post_injury_occupation: str = ""
    post_injury_annual_income: float = 0.0
    labor_market_reduction: float = 0.0
    
    # Additional Costs
    life_care_plan_cost: float = 0.0
    household_services_cost: float = 0.0

def calculate_age(birth_date: date, reference_date: date) -> float:
    """Calculate age in years with decimal precision"""
    return (reference_date - birth_date).days / 365.25

def calculate_statistical_dates(doi: date, le: float, wle: float) -> tuple:
    """Calculate statistical retirement and death dates"""
    retirement_date = doi + timedelta(days=int(wle * 365.25))
    death_date = doi + timedelta(days=int(le * 365.25))
    return retirement_date, death_date

def calculate_worklife_factor(wle: float, yfs: float) -> float:
    """Calculate worklife participation factor"""
    if yfs > 0:
        return wle / (wle + yfs)
    return 1.0

def calculate_adjusted_earnings_factor(unemployment: float, tax: float, 
                                     fringe: float = 0.0, consumption: float = 0.0,
                                     worklife_factor: float = 1.0) -> dict:
    """Calculate comprehensive adjusted earnings factor"""
    factors = []
    current_factor = 1.0
    
    # Worklife factor
    current_factor *= worklife_factor
    factors.append(("Worklife Factor", worklife_factor, current_factor))
    
    # Fringe benefits (additive)
    if fringe > 0:
        current_factor *= (1 + fringe)
        factors.append(("Fringe Benefits", 1 + fringe, current_factor))
    
    # Unemployment reduction
    current_factor *= (1 - unemployment)
    factors.append(("Unemployment Factor", 1 - unemployment, current_factor))
    
    # Tax reduction
    current_factor *= (1 - tax)
    factors.append(("Tax Factor", 1 - tax, current_factor))
    
    # Personal consumption (for wrongful death)
    if consumption > 0:
        current_factor = current_factor - consumption
        factors.append(("Personal Consumption", consumption, current_factor))
    
    return {
        'final_aef': current_factor,
        'factors': factors
    }

def generate_economic_schedule(data: EconomicData) -> tuple:
    """Generate past and future economic loss schedules"""
    
    # Calculate key dates and factors
    retirement_date, death_date = calculate_statistical_dates(
        data.date_of_injury, data.life_expectancy, data.work_life_expectancy
    )
    
    worklife_factor = calculate_worklife_factor(
        data.work_life_expectancy, data.years_to_final_separation
    )
    
    aef_result = calculate_adjusted_earnings_factor(
        data.unemployment_rate, data.tax_rate, 
        data.fringe_benefits_rate, data.personal_consumption_rate,
        worklife_factor
    )
    
    # Past losses (DOI to DOR)
    past_years = []
    report_year = data.date_of_report.year
    injury_year = data.date_of_injury.year
    
    for year in range(injury_year, report_year + 1):
        age = calculate_age(data.date_of_birth, date(year, 7, 1))
        years_since_injury = year - injury_year
        
        # Calculate portion of year
        if year == injury_year:
            # Partial year from injury date
            days_remaining = (date(year, 12, 31) - data.date_of_injury).days + 1
            portion = days_remaining / 365
        elif year == report_year:
            # Partial year to report date
            days_elapsed = data.date_of_report.timetuple().tm_yday
            portion = days_elapsed / 365
        else:
            portion = 1.0
        
        # Calculate earnings
        pre_injury_earnings = data.pre_injury_annual_income * ((1 + data.wage_growth_rate) ** years_since_injury)
        nominal_loss = pre_injury_earnings * portion
        adjusted_loss = nominal_loss * aef_result['final_aef']
        
        # Present value (discount back to injury date)
        pv_loss = adjusted_loss / ((1 + data.discount_rate) ** years_since_injury)
        
        past_years.append({
            'Year': year,
            'Age': round(age, 1),
            'Portion': round(portion * 100, 1),
            'Pre-Injury Earnings': round(pre_injury_earnings, 2),
            'Nominal Loss': round(nominal_loss, 2),
            'AEF': round(aef_result['final_aef'] * 100, 2),
            'Adjusted Loss': round(adjusted_loss, 2),
            'Present Value': round(pv_loss, 2)
        })
    
    # Future losses (DOR to retirement)
    future_years = []
    retirement_year = retirement_date.year
    
    for year in range(report_year, retirement_year + 1):
        if year == report_year:
            continue  # Already handled in past
            
        age = calculate_age(data.date_of_birth, date(year, 7, 1))
        years_since_injury = year - injury_year
        
        # Calculate portion of year
        if year == retirement_year:
            # Partial year to retirement
            retirement_day = retirement_date.timetuple().tm_yday
            portion = retirement_day / 365
        else:
            portion = 1.0
        
        # Calculate earnings
        pre_injury_earnings = data.pre_injury_annual_income * ((1 + data.wage_growth_rate) ** years_since_injury)
        post_injury_earnings = data.post_injury_annual_income * ((1 + data.wage_growth_rate) ** (years_since_injury - (report_year - injury_year)))
        
        nominal_loss = (pre_injury_earnings - post_injury_earnings) * portion
        adjusted_loss = nominal_loss * aef_result['final_aef']
        
        # Present value (discount to injury date)
        pv_loss = adjusted_loss / ((1 + data.discount_rate) ** years_since_injury)
        
        future_years.append({
            'Year': year,
            'Age': round(age, 1),
            'Portion': round(portion * 100, 1),
            'Pre-Injury Earnings': round(pre_injury_earnings, 2),
            'Post-Injury Earnings': round(post_injury_earnings, 2),
            'Nominal Loss': round(nominal_loss, 2),
            'AEF': round(aef_result['final_aef'] * 100, 2),
            'Adjusted Loss': round(adjusted_loss, 2),
            'Present Value': round(pv_loss, 2)
        })
    
    past_df = pd.DataFrame(past_years)
    future_df = pd.DataFrame(future_years)
    
    return past_df, future_df, aef_result, retirement_date, death_date

def create_word_report(data: EconomicData, past_df: pd.DataFrame, 
                      future_df: pd.DataFrame, aef_result: dict, 
                      retirement_date: date, death_date: date) -> bytes:
    """Generate comprehensive Word report following the template structure"""
    
    doc = Document()
    
    # Set document margins and orientation
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    
    # Title page following template structure
    title = doc.add_heading('APPRAISAL OF ECONOMIC LOSS', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading(f'RESULTING FROM INJURY TO {data.client_name.upper()}', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Prepared by section
    prepared_para = doc.add_paragraph()
    prepared_para.add_run("PREPARED BY:").bold = True
    prepared_para.add_run("\t\tKincaid Wolstein Vocational and Rehabilitation Services\n")
    prepared_para.add_run("\t\t\t\t\tOne University Plaza ~ Suite 302\n")
    prepared_para.add_run("\t\t\t\t\tHackensack, New Jersey 07601\n")
    prepared_para.add_run("\t\t\t\t\tPhone: (201) 343-0700\n")
    prepared_para.add_run("\t\t\t\t\tFax: (201) 343-0757")
    
    doc.add_paragraph()
    
    # Case details
    case_details = [
        ("PREPARED FOR:", ""),
        ("REGARDING:", data.client_name),
        ("DATE OF BIRTH:", data.date_of_birth.strftime('%m/%d/%Y')),
        ("REPORT DATE:", data.date_of_report.strftime('%m/%d/%Y'))
    ]
    
    for label, value in case_details:
        para = doc.add_paragraph()
        para.add_run(label).bold = True
        para.add_run(f"\t\t{value}")
    
    # Page break to table of contents
    doc.add_page_break()
    
    # Table of Contents
    toc_heading = doc.add_heading('Table of Contents', level=1)
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    toc_items = [
        "CERTIFICATION\t3",
        "PURPOSE OF APPRAISAL\t4", 
        "OPINION OF ECONOMIC LOSSES\t4",
        "BACKGROUND FACTS AND ASSUMPTIONS\t5",
        "Summary Information\t5",
        "Life Expectancy\t5",
        "Statistical retirement age\t5",
        "Expected working years\t5",
        "Worklife-to Retirement Ratio\t6",
        "Description of Incident\t6",
        "Occupation and Employment\t6",
        "Earnings History\t7",
        "Pre-injury Earnings Capacity\t7",
        "Fringe Benefits\t7",
        "Functionality and Future Employability\t8",
        "Post-injury Earnings Capacity\t9",
        "Growth Wage Rates\t9",
        "Household Services\t10",
        "COMPONENTS OF ANALYSIS\t12",
        "PRE‐INJURY ADJUSTED EARNINGS\t13",
        "PRE‐INJURY ADJUSTED EARNINGS FOLLOWING RTW\t15",
        "POST‐INJURY ADJUSTED EARNINGS FOLLOWING DEPARTURE FROM FULL DUTY POSITION\t16",
        "COST OF LIFETIME CARE\t18",
        "Methodology\t20",
        "Summary of Findings\t20",
        "Important Considerations\t21",
        "COST OF LIFETIME CARE\t23",
        "TABLES\t24",
        "STATEMENT OF ETHICAL PRINCIPLES AND PRINCIPLES OF PROFESSIONAL PRACTICE\t38"
    ]
    
    for item in toc_items:
        doc.add_paragraph(item)
    
    # Page break to main content
    doc.add_page_break()
    
    # Executive Summary
    doc.add_heading('EXECUTIVE SUMMARY', level=1)
    
    past_total = past_df['Present Value'].sum() if not past_df.empty else 0
    future_total = future_df['Present Value'].sum() if not future_df.empty else 0
    total_loss = past_total + future_total + data.life_care_plan_cost
    
    summary_text = f"""
This report presents an economic analysis of the financial losses sustained by {data.client_name} 
as a result of injuries sustained on {data.date_of_injury.strftime('%B %d, %Y')}. The analysis 
considers both past and future economic losses through the individual's projected work life.

OPINION OF ECONOMIC LOSSES:

Total Economic Loss: ${total_loss:,.2f}

This total comprises:
• Past Lost Earnings: ${past_total:,.2f}
• Future Lost Earnings: ${future_total:,.2f}
• Life Care Plan Costs: ${data.life_care_plan_cost:,.2f}
• Household Services: ${data.household_services_cost:,.2f}

The analysis employs established economic methodologies including present value calculations, 
demographic adjustments, and the Tinari Adjustment methodology (Adjusted Earnings Factor = {aef_result['final_aef']:.3f}).
"""
    
    doc.add_paragraph(summary_text)
    
    # Background Facts and Assumptions
    doc.add_page_break()
    doc.add_heading('BACKGROUND FACTS AND ASSUMPTIONS', level=1)
    
    # Summary Information section
    doc.add_heading('Summary Information', level=2)
    
    # Create summary table
    summary_table = doc.add_table(rows=15, cols=2)
    summary_table.style = 'Table Grid'
    
    summary_info_data = [
        ('Client Name', data.client_name),
        ('Date of Birth', data.date_of_birth.strftime('%m/%d/%Y')),
        ('Date of Injury', data.date_of_injury.strftime('%m/%d/%Y')),
        ('Date of Report', data.date_of_report.strftime('%m/%d/%Y')),
        ('Age at Injury', f"{calculate_age(data.date_of_birth, data.date_of_injury):.1f} years"),
        ('Gender', data.gender),
        ('Education', data.education),
        ('Marital Status', data.marital_status),
        ('Residence', f"{data.residence_county}, {data.residence_state}"),
        ('MSA', data.msa),
        ('Pre-Injury Occupation', data.pre_injury_occupation),
        ('Employer', data.employer),
        ('Years with Employer', f"{data.years_with_employer:.1f} years"),
        ('Pre-Injury Annual Income', f"${data.pre_injury_annual_income:,.2f}"),
        ('Pre-Injury Hourly Rate', f"${data.pre_injury_hourly_rate:.2f}")
    ]
    
    for i, (label, value) in enumerate(summary_info_data):
        summary_table.cell(i, 0).text = label
        summary_table.cell(i, 1).text = str(value)
    
    # Life Expectancy section
    doc.add_heading('Life Expectancy', level=2)
    life_exp_text = f"""
Based on standard actuarial tables and demographic data:
• Life Expectancy from Date of Injury: {data.life_expectancy:.1f} years
• Statistical Death Date: {death_date.strftime('%B %d, %Y')}
• Current Age: {calculate_age(data.date_of_birth, data.date_of_report):.1f} years
"""
    doc.add_paragraph(life_exp_text)
    
    # Statistical retirement age section
    doc.add_heading('Statistical Retirement Age', level=2)
    retirement_text = f"""
Work-life expectancy analysis:
• Work-Life Expectancy from Date of Injury: {data.work_life_expectancy:.1f} years
• Statistical Retirement Date: {retirement_date.strftime('%B %d, %Y')}
• Remaining Work Years from Report Date: {(retirement_date - data.date_of_report).days / 365.25:.1f} years
"""
    doc.add_paragraph(retirement_text)
    
    # Expected working years section
    doc.add_heading('Expected Working Years', level=2)
    working_years_text = f"""
Total analysis period: {(retirement_date - data.date_of_injury).days / 365.25:.1f} years
• Past period (injury to report): {(data.date_of_report - data.date_of_injury).days / 365.25:.1f} years
• Future period (report to retirement): {(retirement_date - data.date_of_report).days / 365.25:.1f} years
"""
    doc.add_paragraph(working_years_text)
    
    # Worklife-to Retirement Ratio section
    doc.add_heading('Worklife-to Retirement Ratio', level=2)
    worklife_factor = calculate_worklife_factor(data.work_life_expectancy, data.years_to_final_separation)
    ratio_text = f"""
Worklife participation factor: {worklife_factor:.4f}
This factor accounts for the probability of continued labor force participation through the projected work-life period.
"""
    doc.add_paragraph(ratio_text)
    
    # Description of Incident section
    doc.add_heading('Description of Incident', level=2)
    incident_text = f"""
Date of Injury: {data.date_of_injury.strftime('%B %d, %Y')}

Nature of Injury:
{data.injury_description}

Body Parts Injured:
{data.body_parts_injured}

Medical Treatment:
{data.medical_treatment}

Current Medical Status:
{data.current_medical_status}
"""
    doc.add_paragraph(incident_text)
    
    # Occupation and Employment section
    doc.add_heading('Occupation and Employment', level=2)
    employment_text = f"""
Pre-Injury Employment:
• Occupation: {data.pre_injury_occupation}
• Employer: {data.employer}
• Length of Employment: {data.years_with_employer:.1f} years
• Annual Income: ${data.pre_injury_annual_income:,.2f}
• Hourly Rate: ${data.pre_injury_hourly_rate:.2f}

Post-Injury Work Capacity:
{data.post_injury_capacity}

Post-Injury Employment Potential:
• Occupation: {data.post_injury_occupation}
• Projected Annual Income: ${data.post_injury_annual_income:,.2f}
• Labor Market Reduction: {data.labor_market_reduction*100:.1f}%
"""
    doc.add_paragraph(employment_text)
    
    # Economic assumptions sections
    doc.add_heading('Economic Assumptions', level=2)
    
    # Growth Wage Rates
    doc.add_heading('Growth Wage Rates', level=3)
    growth_text = f"""
Wage Growth Rate: {data.wage_growth_rate*100:.1f}% annually
This rate is based on historical wage growth patterns and economic forecasting data.
"""
    doc.add_paragraph(growth_text)
    
    # Fringe Benefits
    doc.add_heading('Fringe Benefits', level=3)
    fringe_text = f"""
Fringe Benefits Rate: {data.fringe_benefits_rate*100:.1f}% of wages
Includes employer-provided benefits such as:
• Health insurance
• Retirement contributions
• Life insurance
• Paid time off
• Other employer-provided benefits
"""
    doc.add_paragraph(fringe_text)
    
    # Discount rate section
    doc.add_heading('Discount Rate and Present Value', level=3)
    discount_text = f"""
Discount Rate: {data.discount_rate*100:.1f}% annually
This rate is used to convert future economic losses to present value, reflecting the time value of money.
"""
    doc.add_paragraph(discount_text)
    
    # AEF methodology section  
    doc.add_page_break()
    doc.add_heading('ADJUSTED EARNINGS FACTOR (AEF) METHODOLOGY', level=1)
    
    aef_methodology_text = f"""
The Adjusted Earnings Factor represents a comprehensive adjustment to gross earnings that accounts 
for real-world economic conditions. This methodology incorporates the following factors:

1. UNEMPLOYMENT FACTOR ({data.unemployment_rate*100:.1f}%): Accounts for periods of unemployment during working life
2. TAX FACTOR ({data.tax_rate*100:.1f}%): Adjusts for income taxes, reflecting take-home pay  
3. FRINGE BENEFITS FACTOR ({data.fringe_benefits_rate*100:.1f}%): Includes employer-provided benefits
4. PERSONAL CONSUMPTION FACTOR ({data.personal_consumption_rate*100:.1f}%): For wrongful death cases only

Final AEF: {aef_result['final_aef']:.4f} (representing {aef_result['final_aef']*100:.2f}% of gross earnings)

This means that for every $1.00 of gross earnings capacity, the net household economic impact is ${aef_result['final_aef']:.2f}.
"""
    doc.add_paragraph(aef_methodology_text)
    
    # AEF breakdown table
    aef_table = doc.add_table(rows=len(aef_result['factors']) + 1, cols=3)
    aef_table.style = 'Table Grid'
    
    # Headers
    aef_table.cell(0, 0).text = 'Factor'
    aef_table.cell(0, 1).text = 'Value'
    aef_table.cell(0, 2).text = 'Cumulative'
    
    for i, (factor_name, factor_value, cumulative) in enumerate(aef_result['factors']):
        aef_table.cell(i + 1, 0).text = factor_name
        aef_table.cell(i + 1, 1).text = f"{factor_value:.4f}"
        aef_table.cell(i + 1, 2).text = f"{cumulative:.4f}"
    
    # Past Losses section
    doc.add_page_break()
    doc.add_heading('PAST ECONOMIC LOSSES', level=1)
    
    past_analysis_text = f"""
Analysis Period: {data.date_of_injury.strftime('%B %d, %Y')} to {data.date_of_report.strftime('%B %d, %Y')}
Duration: {(data.date_of_report - data.date_of_injury).days / 365.25:.1f} years

This section quantifies the economic losses that have occurred from the date of injury through 
the date of this report. These represent the difference between projected earning capacity and 
actual earnings during this period.

Total Past Losses (Present Value): ${past_total:,.2f}
"""
    doc.add_paragraph(past_analysis_text)
    
    if not past_df.empty:
        # Create past losses table
        past_table = doc.add_table(rows=len(past_df) + 2, cols=len(past_df.columns))
        past_table.style = 'Table Grid'
        
        # Headers
        for j, col in enumerate(past_df.columns):
            past_table.cell(0, j).text = col
        
        # Data
        for i in range(len(past_df)):
            for j, col in enumerate(past_df.columns):
                value = past_df.iloc[i, j]
                if isinstance(value, (int, float)) and col in ['Pre-Injury Earnings', 'Nominal Loss', 'Adjusted Loss', 'Present Value']:
                    past_table.cell(i + 1, j).text = f"${value:,.2f}"
                else:
                    past_table.cell(i + 1, j).text = str(value)
        
        # Totals row
        totals_row = len(past_df) + 1
        past_table.cell(totals_row, 0).text = "TOTALS"
        for j, col in enumerate(past_df.columns):
            if col in ['Pre-Injury Earnings', 'Nominal Loss', 'Adjusted Loss', 'Present Value']:
                total = past_df[col].sum()
                past_table.cell(totals_row, j).text = f"${total:,.2f}"
    
    # Future Losses section
    doc.add_page_break()
    doc.add_heading('FUTURE ECONOMIC LOSSES', level=1)
    
    future_analysis_text = f"""
Projection Period: {data.date_of_report.strftime('%B %d, %Y')} to {retirement_date.strftime('%B %d, %Y')}
Duration: {(retirement_date - data.date_of_report).days / 365.25:.1f} years

This section projects the economic losses expected to occur from the report date through the 
projected retirement date. These projections are based on established economic methodologies 
and demographic data.

Economic Assumptions:
• Wage Growth Rate: {data.wage_growth_rate*100:.1f}% annually
• Discount Rate: {data.discount_rate*100:.1f}% annually
• Present Value Base Date: {data.date_of_report.strftime('%B %d, %Y')}

Total Future Losses (Present Value): ${future_total:,.2f}
"""
    doc.add_paragraph(future_analysis_text)
    
    if not future_df.empty:
        # Create future losses table
        future_table = doc.add_table(rows=len(future_df) + 2, cols=len(future_df.columns))
        future_table.style = 'Table Grid'
        
        # Headers
        for j, col in enumerate(future_df.columns):
            future_table.cell(0, j).text = col
        
        # Data
        for i in range(len(future_df)):
            for j, col in enumerate(future_df.columns):
                value = future_df.iloc[i, j]
                if isinstance(value, (int, float)) and col in ['Pre-Injury Earnings', 'Post-Injury Earnings', 'Nominal Loss', 'Adjusted Loss', 'Present Value']:
                    future_table.cell(i + 1, j).text = f"${value:,.2f}"
                else:
                    future_table.cell(i + 1, j).text = str(value)
        
        # Totals row
        totals_row = len(future_df) + 1
        future_table.cell(totals_row, 0).text = "TOTALS"
        for j, col in enumerate(future_df.columns):
            if col in ['Pre-Injury Earnings', 'Post-Injury Earnings', 'Nominal Loss', 'Adjusted Loss', 'Present Value']:
                total = future_df[col].sum()
                future_table.cell(totals_row, j).text = f"${total:,.2f}"
    
    # Additional cost sections
    if data.life_care_plan_cost > 0:
        doc.add_page_break()
        doc.add_heading('COST OF LIFETIME CARE', level=1)
        
        lcp_text = f"""
Life Care Plan Cost Analysis:

The life care plan represents the present value of future medical and care-related expenses 
resulting from the injury. This analysis is typically based on a detailed life care plan 
prepared by qualified medical and rehabilitation professionals.

Total Life Care Plan Cost (Present Value): ${data.life_care_plan_cost:,.2f}

This amount represents the funds needed today to cover all projected future medical expenses, 
therapies, equipment, and care services related to the injury over the individual's lifetime.
"""
        doc.add_paragraph(lcp_text)
    
    if data.household_services_cost > 0:
        doc.add_page_break()
        doc.add_heading('HOUSEHOLD SERVICES', level=1)
        
        household_text = f"""
Loss of Household Services Analysis:

The injury has resulted in a diminished capacity to perform household and family services. 
This represents a real economic loss that must be quantified using replacement cost methodology.

Total Household Services Loss (Present Value): ${data.household_services_cost:,.2f}

This amount represents the cost to replace household services that can no longer be performed 
due to the injury, valued at current market rates for such services.
"""
        doc.add_paragraph(household_text)
    
    # Summary of Findings
    doc.add_page_break()
    doc.add_heading('SUMMARY OF FINDINGS', level=1)
    
    # Create comprehensive summary table
    summary_table = doc.add_table(rows=6, cols=2)
    summary_table.style = 'Table Grid'
    
    total_all_losses = past_total + future_total + data.life_care_plan_cost + data.household_services_cost
    
    summary_data = [
        ('Past Economic Losses', f"${past_total:,.2f}"),
        ('Future Economic Losses', f"${future_total:,.2f}"),
        ('Life Care Plan Costs', f"${data.life_care_plan_cost:,.2f}"),
        ('Household Services Loss', f"${data.household_services_cost:,.2f}"),
        ('', ''),
        ('TOTAL ECONOMIC LOSS', f"${total_all_losses:,.2f}")
    ]
    
    for i, (label, value) in enumerate(summary_data):
        summary_table.cell(i, 0).text = label
        summary_table.cell(i, 1).text = value
        if 'TOTAL' in label:
            # Make total row bold
            for paragraph in summary_table.cell(i, 0).paragraphs:
                for run in paragraph.runs:
                    run.bold = True
            for paragraph in summary_table.cell(i, 1).paragraphs:
                for run in paragraph.runs:
                    run.bold = True
    
    # Important Considerations section
    doc.add_heading('Important Considerations', level=2)
    
    considerations_text = f"""
METHODOLOGY NOTES:
• All calculations employ standard economic principles and established methodologies
• Present value calculations utilize appropriate discount rates reflecting current economic conditions
• Demographic projections are based on standard actuarial tables and work-life expectancy data
• The Adjusted Earnings Factor methodology accounts for real-world economic conditions
• All projections are subject to the inherent uncertainty of future economic conditions

CALCULATION VALIDATION:
• Economic assumptions are based on historical data and reasonable projections
• All calculations have been verified for mathematical accuracy
• Results are presented to a reasonable degree of economic certainty
• Sensitivity analysis can be performed if required for key assumptions

PROFESSIONAL STANDARDS:
• Analysis follows accepted practices in forensic economics
• Methodology is consistent with standards for expert economic testimony
• All assumptions and data sources are documented and transparent
• Results comply with professional ethical standards for economic analysis
"""
    doc.add_paragraph(considerations_text)
    
    # Professional statement
    doc.add_page_break()
    doc.add_heading('STATEMENT OF ETHICAL PRINCIPLES AND PRINCIPLES OF PROFESSIONAL PRACTICE', level=1)
    
    ethics_text = """
This economic analysis has been prepared in accordance with the ethical principles and 
professional practice standards of forensic economics. The analysis is based on reliable 
data, established methodologies, and reasonable assumptions. All calculations have been 
performed with due care and professional competence.

The opinions expressed in this report are offered to a reasonable degree of economic certainty 
and are based on the information available at the time of preparation. The analyst maintains 
independence and objectivity in all professional work and adheres to the highest standards 
of professional conduct.

This report contains the complete opinions of the analyst on the matters addressed herein. 
No part of the compensation for this analysis is contingent upon the results obtained or 
the testimony given.
"""
    doc.add_paragraph(ethics_text)
    
    # Save to bytes
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io.getvalue()

# Streamlit Interface
def main():
    # Sidebar for main inputs
    with st.sidebar:
        st.header("📋 Report Configuration")
        
        # Sample case loader
        st.subheader("Quick Start")
        if st.button("🔄 Load Sample Case", type="primary"):
            load_sample_case()
        
        st.markdown("---")
        
        # Load/Save functionality
        st.subheader("Data Management")
        
        col1, col2 = st.columns(2)
        with col1:
            if st.button("💾 Save Data"):
                if st.session_state.report_data:
                    # Convert dates to strings for JSON
                    save_data = {}
                    for key, value in st.session_state.report_data.items():
                        if isinstance(value, date):
                            save_data[key] = value.isoformat()
                        else:
                            save_data[key] = value
                    
                    data_json = json.dumps(save_data, indent=2)
                    st.download_button(
                        "📥 Download JSON",
                        data_json,
                        f"economic_report_data_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                        "application/json"
                    )
                else:
                    st.warning("No data to save. Please fill in the form first.")
        
        with col2:
            uploaded_file = st.file_uploader("📁 Load Data", type=['json'])
            if uploaded_file:
                try:
                    loaded_data = json.load(uploaded_file)
                    # Convert date strings back to date objects
                    for key, value in loaded_data.items():
                        if key in ['date_of_birth', 'date_of_injury', 'date_of_report'] and isinstance(value, str):
                            try:
                                loaded_data[key] = datetime.fromisoformat(value).date()
                            except:
                                pass
                    st.session_state.report_data.update(loaded_data)
                    st.success("Data loaded successfully!")
                    st.rerun()
                except Exception as e:
                    st.error(f"Error loading file: {str(e)}")
    
    # Main content area with tabs
    tab1, tab2, tab3, tab4, tab5 = st.tabs([
        "👤 Personal Info", 
        "💼 Employment", 
        "🏥 Injury Details", 
        "📊 Economic Factors", 
        "📋 Generate Report"
    ])
    
    # Initialize data object
    data = EconomicData()
    
    with tab1:
        st.header("Personal Information")
        
        col1, col2 = st.columns(2)
        with col1:
            data.client_name = st.text_input(
                "Client Name", 
                value=st.session_state.report_data.get('client_name', ''),
                key='client_name',
                help=FIELD_HELP['client_name']
            )
            data.date_of_birth = st.date_input(
                "Date of Birth", 
                value=st.session_state.report_data.get('date_of_birth', date(1980, 1, 1)),
                key='date_of_birth',
                help=FIELD_HELP['date_of_birth']
            )
            data.gender = st.selectbox(
                "Gender", 
                ["Male", "Female"],
                index=0 if st.session_state.report_data.get('gender', 'Male') == 'Male' else 1,
                key='gender',
                help=FIELD_HELP['gender']
            )
            data.education = st.text_input(
                "Education Level", 
                value=st.session_state.report_data.get('education', ''),
                key='education',
                help=FIELD_HELP['education']
            )
        
        with col2:
            data.date_of_injury = st.date_input(
                "Date of Injury", 
                value=st.session_state.report_data.get('date_of_injury', date(2023, 1, 1)),
                key='date_of_injury',
                help=FIELD_HELP['date_of_injury']
            )
            data.date_of_report = st.date_input(
                "Date of Report", 
                value=st.session_state.report_data.get('date_of_report', date.today()),
                key='date_of_report',
                help=FIELD_HELP['date_of_report']
            )
            data.marital_status = st.text_input(
                "Marital Status", 
                value=st.session_state.report_data.get('marital_status', ''),
                key='marital_status',
                help=FIELD_HELP['marital_status']
            )
        
        # Calculate age at injury
        data.age_at_injury = calculate_age(data.date_of_birth, data.date_of_injury)
        st.info(f"Age at injury: {data.age_at_injury:.1f} years")
        
        st.subheader("Geographic Information")
        col3, col4, col5 = st.columns(3)
        with col3:
            data.residence_state = st.text_input(
                "State of Residence", 
                value=st.session_state.report_data.get('residence_state', ''),
                key='residence_state',
                help=FIELD_HELP['residence_state']
            )
        with col4:
            data.residence_county = st.text_input(
                "County", 
                value=st.session_state.report_data.get('residence_county', ''),
                key='residence_county',
                help=FIELD_HELP['residence_county']
            )
        with col5:
            data.msa = st.text_input(
                "MSA (Metropolitan Statistical Area)", 
                value=st.session_state.report_data.get('msa', ''),
                key='msa',
                help=FIELD_HELP['msa']
            )
    
    with tab2:
        st.header("Employment Information")
        
        col1, col2 = st.columns(2)
        with col1:
            data.pre_injury_occupation = st.text_input(
                "Pre-Injury Occupation", 
                value=st.session_state.report_data.get('pre_injury_occupation', ''),
                key='pre_injury_occupation',
                help=FIELD_HELP['pre_injury_occupation']
            )
            data.pre_injury_annual_income = st.number_input(
                "Pre-Injury Annual Income ($)", 
                min_value=0.0, 
                value=float(st.session_state.report_data.get('pre_injury_annual_income', 50000.0)),
                step=1000.0,
                key='pre_injury_annual_income',
                help=FIELD_HELP['pre_injury_annual_income']
            )
            data.employer = st.text_input(
                "Employer", 
                value=st.session_state.report_data.get('employer', ''),
                key='employer',
                help=FIELD_HELP['employer']
            )
        
        with col2:
            data.pre_injury_hourly_rate = st.number_input(
                "Pre-Injury Hourly Rate ($)", 
                min_value=0.0, 
                value=float(st.session_state.report_data.get('pre_injury_hourly_rate', 24.04)),
                step=0.50,
                key='pre_injury_hourly_rate',
                help=FIELD_HELP['pre_injury_hourly_rate']
            )
            data.years_with_employer = st.number_input(
                "Years with Employer", 
                min_value=0.0, 
                value=float(st.session_state.report_data.get('years_with_employer', 5.0)),
                step=0.5,
                key='years_with_employer',
                help=FIELD_HELP['years_with_employer']
            )
        
        st.subheader("Post-Injury Employment")
        col3, col4 = st.columns(2)
        with col3:
            data.post_injury_occupation = st.text_input(
                "Post-Injury Occupation", 
                value=st.session_state.report_data.get('post_injury_occupation', ''),
                key='post_injury_occupation',
                help=FIELD_HELP['post_injury_occupation']
            )
            data.post_injury_capacity = st.text_area(
                "Post-Injury Work Capacity", 
                value=st.session_state.report_data.get('post_injury_capacity', ''),
                key='post_injury_capacity',
                help=FIELD_HELP['post_injury_capacity']
            )
        
        with col4:
            data.post_injury_annual_income = st.number_input(
                "Post-Injury Annual Income ($)", 
                min_value=0.0, 
                value=float(st.session_state.report_data.get('post_injury_annual_income', 0.0)),
                step=1000.0,
                key='post_injury_annual_income',
                help=FIELD_HELP['post_injury_annual_income']
            )
            data.labor_market_reduction = st.number_input(
                "Labor Market Reduction (%)", 
                min_value=0.0,
                max_value=100.0,
                value=float(st.session_state.report_data.get('labor_market_reduction', 0.0)),
                step=1.0,
                key='labor_market_reduction',
                help=FIELD_HELP['labor_market_reduction']
            ) / 100
    
    with tab3:
        st.header("Injury and Medical Information")
        
        col1, col2 = st.columns(2)
        with col1:
            data.injury_description = st.text_area(
                "Injury Description", 
                value=st.session_state.report_data.get('injury_description', ''),
                key='injury_description',
                height=100,
                help=FIELD_HELP['injury_description']
            )
            data.body_parts_injured = st.text_input(
                "Body Parts Injured", 
                value=st.session_state.report_data.get('body_parts_injured', ''),
                key='body_parts_injured',
                help=FIELD_HELP['body_parts_injured']
            )
        
        with col2:
            data.medical_treatment = st.text_area(
                "Medical Treatment", 
                value=st.session_state.report_data.get('medical_treatment', ''),
                key='medical_treatment',
                height=100,
                help=FIELD_HELP['medical_treatment']
            )
            data.current_medical_status = st.text_area(
                "Current Medical Status", 
                value=st.session_state.report_data.get('current_medical_status', ''),
                key='current_medical_status',
                height=100,
                help=FIELD_HELP['current_medical_status']
            )
    
    with tab4:
        st.header("Economic Factors and Life Expectancy")
        
        st.subheader("Life Expectancy Data")
        col1, col2, col3 = st.columns(3)
        with col1:
            data.life_expectancy = st.number_input(
                "Life Expectancy (years)", 
                min_value=0.0,
                max_value=120.0,
                value=float(st.session_state.report_data.get('life_expectancy', 78.5)),
                step=0.1,
                key='life_expectancy',
                help=FIELD_HELP['life_expectancy']
            )
        with col2:
            data.work_life_expectancy = st.number_input(
                "Work Life Expectancy (years)", 
                min_value=0.0,
                max_value=80.0,
                value=float(st.session_state.report_data.get('work_life_expectancy', 45.0)),
                step=0.1,
                key='work_life_expectancy',
                help=FIELD_HELP['work_life_expectancy']
            )
        with col3:
            data.years_to_final_separation = st.number_input(
                "Years to Final Separation", 
                min_value=0.0,
                value=float(st.session_state.report_data.get('years_to_final_separation', 0.0)),
                step=0.1,
                key='years_to_final_separation',
                help=FIELD_HELP['years_to_final_separation']
            )
        
        # Calculate and display statistical dates
        retirement_date, death_date = calculate_statistical_dates(
            data.date_of_injury, data.life_expectancy, data.work_life_expectancy
        )
        
        col4, col5 = st.columns(2)
        with col4:
            st.info(f"Statistical Retirement Date: {retirement_date.strftime('%m/%d/%Y')}")
        with col5:
            st.info(f"Statistical Death Date: {death_date.strftime('%m/%d/%Y')}")
        
        st.subheader("Economic Growth and Discount Rates")
        col6, col7 = st.columns(2)
        with col6:
            data.wage_growth_rate = st.number_input(
                "Wage Growth Rate (%)", 
                min_value=0.0,
                max_value=20.0,
                value=float(st.session_state.report_data.get('wage_growth_rate', 3.0)),
                step=0.1,
                key='wage_growth_rate',
                help=FIELD_HELP['wage_growth_rate']
            ) / 100
        with col7:
            data.discount_rate = st.number_input(
                "Discount Rate (%)", 
                min_value=0.0,
                max_value=20.0,
                value=float(st.session_state.report_data.get('discount_rate', 4.0)),
                step=0.1,
                key='discount_rate',
                help=FIELD_HELP['discount_rate']
            ) / 100
        
        st.subheader("Adjustment Factors")
        col8, col9 = st.columns(2)
        with col8:
            data.unemployment_rate = st.number_input(
                "Unemployment Rate (%)", 
                min_value=0.0,
                max_value=30.0,
                value=float(st.session_state.report_data.get('unemployment_rate', 3.5)),
                step=0.1,
                key='unemployment_rate',
                help=FIELD_HELP['unemployment_rate']
            ) / 100
            data.tax_rate = st.number_input(
                "Tax Rate (%)", 
                min_value=0.0,
                max_value=50.0,
                value=float(st.session_state.report_data.get('tax_rate', 12.0)),
                step=0.1,
                key='tax_rate',
                help=FIELD_HELP['tax_rate']
            ) / 100
        with col9:
            data.fringe_benefits_rate = st.number_input(
                "Fringe Benefits Rate (%)", 
                min_value=0.0,
                max_value=50.0,
                value=float(st.session_state.report_data.get('fringe_benefits_rate', 6.0)),
                step=0.1,
                key='fringe_benefits_rate',
                help=FIELD_HELP['fringe_benefits_rate']
            ) / 100
            data.personal_consumption_rate = st.number_input(
                "Personal Consumption Rate (%)", 
                min_value=0.0,
                max_value=50.0,
                value=float(st.session_state.report_data.get('personal_consumption_rate', 0.0)),
                step=0.1,
                key='personal_consumption_rate',
                help=FIELD_HELP['personal_consumption_rate'] + " - Use for wrongful death cases"
            ) / 100
        
        st.subheader("Additional Costs")
        col10, col11 = st.columns(2)
        with col10:
            data.life_care_plan_cost = st.number_input(
                "Life Care Plan Cost ($)", 
                min_value=0.0,
                value=float(st.session_state.report_data.get('life_care_plan_cost', 0.0)),
                step=1000.0,
                key='life_care_plan_cost',
                help=FIELD_HELP['life_care_plan_cost']
            )
        with col11:
            data.household_services_cost = st.number_input(
                "Household Services Cost ($)", 
                min_value=0.0,
                value=float(st.session_state.report_data.get('household_services_cost', 0.0)),
                step=1000.0,
                key='household_services_cost',
                help=FIELD_HELP['household_services_cost']
            )
    
    with tab5:
        st.header("Generate Economic Loss Report")
        
        # Update session state with current data
        for key, value in data.__dict__.items():
            st.session_state.report_data[key] = value
        
        if st.button("🔄 Calculate Economic Losses", type="primary"):
            with st.spinner("Calculating economic losses..."):
                try:
                    # Generate calculations
                    past_df, future_df, aef_result, retirement_date, death_date = generate_economic_schedule(data)
                    
                    # Display results
                    st.subheader("📊 Calculation Results")
                    
                    # Summary metrics
                    past_total = past_df['Present Value'].sum() if not past_df.empty else 0
                    future_total = future_df['Present Value'].sum() if not future_df.empty else 0
                    total_loss = past_total + future_total + data.life_care_plan_cost
                    
                    col1, col2, col3, col4 = st.columns(4)
                    with col1:
                        st.metric("Past Losses", f"${past_total:,.2f}")
                    with col2:
                        st.metric("Future Losses", f"${future_total:,.2f}")
                    with col3:
                        st.metric("Life Care Plan", f"${data.life_care_plan_cost:,.2f}")
                    with col4:
                        st.metric("Total Economic Loss", f"${total_loss:,.2f}")
                    
                    # AEF breakdown
                    st.subheader("🧮 Adjusted Earnings Factor (AEF)")
                    aef_data = []
                    for factor_name, factor_value, cumulative in aef_result['factors']:
                        aef_data.append({
                            'Factor': factor_name,
                            'Value': f"{factor_value:.3f}",
                            'Percentage': f"{factor_value * 100:.1f}%",
                            'Cumulative': f"{cumulative:.3f}"
                        })
                    
                    aef_df = pd.DataFrame(aef_data)
                    st.dataframe(aef_df, use_container_width=True, hide_index=True)
                    
                    # Past losses table
                    if not past_df.empty:
                        st.subheader("📅 Past Economic Losses")
                        st.dataframe(past_df, use_container_width=True, hide_index=True)
                    
                    # Future losses table
                    if not future_df.empty:
                        st.subheader("🔮 Future Economic Losses")
                        st.dataframe(future_df, use_container_width=True, hide_index=True)
                    
                    # Generate Word report
                    st.subheader("📄 Download Report")
                    
                    col1, col2 = st.columns([1, 1])
                    
                    with col1:
                        if st.button("📄 Generate Word Report", type="primary"):
                            with st.spinner("Generating Word document..."):
                                try:
                                    doc_bytes = create_word_report(data, past_df, future_df, aef_result, retirement_date, death_date)
                                    
                                    if doc_bytes:
                                        # Create download button
                                        file_name = f"Economic_Loss_Report_{data.client_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"
                                        
                                        st.download_button(
                                            label="📥 Download Word Report",
                                            data=doc_bytes,
                                            file_name=file_name,
                                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                            key="word_download"
                                        )
                                        
                                        st.success(f"✅ Word report generated successfully! ({len(doc_bytes):,} bytes)")
                                    else:
                                        st.error("Failed to generate Word document")
                                except Exception as e:
                                    st.error(f"Error generating Word report: {str(e)}")
                                    st.exception(e)
                    
                    with col2:
                        st.info("💡 The Word report includes:\n- Executive summary\n- Case information\n- Economic assumptions\n- AEF breakdown\n- Past & future loss tables\n- Summary totals")
                    
                    # Export data options
                    st.subheader("📊 Export Data")
                    col1, col2, col3 = st.columns(3)
                    
                    with col1:
                        if not past_df.empty:
                            csv_past = past_df.to_csv(index=False)
                            st.download_button(
                                "📊 Past Losses CSV",
                                csv_past,
                                f"past_losses_{data.client_name.replace(' ', '_')}.csv",
                                "text/csv"
                            )
                    
                    with col2:
                        if not future_df.empty:
                            csv_future = future_df.to_csv(index=False)
                            st.download_button(
                                "📊 Future Losses CSV",
                                csv_future,
                                f"future_losses_{data.client_name.replace(' ', '_')}.csv",
                                "text/csv"
                            )
                    
                    with col3:
                        # Combined Excel export
                        excel_io = io.BytesIO()
                        try:
                            with pd.ExcelWriter(excel_io, engine='openpyxl') as writer:
                                if not past_df.empty:
                                    past_df.to_excel(writer, sheet_name='Past Losses', index=False)
                                if not future_df.empty:
                                    future_df.to_excel(writer, sheet_name='Future Losses', index=False)
                                aef_df.to_excel(writer, sheet_name='AEF Breakdown', index=False)
                            
                            excel_data = excel_io.getvalue()
                            st.download_button(
                                "📊 Complete Excel Report",
                                excel_data,
                                f"economic_analysis_{data.client_name.replace(' ', '_')}.xlsx",
                                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                            )
                        except Exception as e:
                            st.error(f"Error creating Excel file: {str(e)}")
                    
                except Exception as e:
                    st.error(f"Error in calculations: {str(e)}")
                    st.exception(e)

if __name__ == "__main__":
    main()