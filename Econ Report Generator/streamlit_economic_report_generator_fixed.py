import streamlit as st
import pandas as pd
import numpy as np
from datetime import date, datetime, timedelta
from dataclasses import dataclass
import io
from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.shared import OxmlElement, qn
import json
import matplotlib.pyplot as plt
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.table import _Cell
import base64

# Import the complete professional template
from complete_professional_template import create_complete_professional_report

# Page configuration
st.set_page_config(
    page_title="Economic Report Generator", 
    layout="wide", 
    initial_sidebar_state="expanded"
)

st.title("📊 Economic Loss Report Generator")
st.markdown("Generate comprehensive economic loss reports with professional calculations and formatting.")

# Sample case data for testing
# Comprehensive sample case for testing all functionality
SAMPLE_CASE = {
    'client_name': 'FIRST LAST',
    'date_of_birth': date(1980, 3, 15),
    'date_of_injury': date(2023, 6, 1),
    'date_of_report': date(2025, 6, 20),
    'gender': 'Male',
    'education': 'Bachelor\'s Degree in Mechanical Engineering',
    'marital_status': 'Married',
    'residence_state': 'New Jersey',
    'residence_county': 'Bergen County',
    'msa': 'New York-Newark-Jersey City, NY-NJ-PA',
    'pre_injury_occupation': 'Mechanical Engineer - Senior Level',
    'pre_injury_annual_income': 85000.0,
    'pre_injury_hourly_rate': 40.87,
    'employer': 'ABC Manufacturing Corporation',
    'years_with_employer': 8.5,
    'injury_description': 'Work-related back injury sustained while lifting heavy machinery components during routine maintenance. Resulted in herniated discs at L4-L5 and chronic pain syndrome limiting physical capacity and mobility.',
    'body_parts_injured': 'Lower lumbar spine, specifically L4-L5 vertebrae with disc herniation',
    'medical_treatment': 'Emergency department treatment, comprehensive MRI and CT diagnostics, intensive physical therapy (6 months), pain management program, epidural steroid injections, ongoing medical monitoring',
    'current_medical_status': 'Maximum medical improvement achieved. Permanent work restrictions: no lifting >20 lbs, limited standing/walking periods, cannot perform fieldwork requiring physical activity',
    'life_expectancy': 78.5,
    'work_life_expectancy': 42.0,
    'years_to_final_separation': 0.0,
    'wage_growth_rate': 0.032,
    'discount_rate': 0.045,
    'unemployment_rate': 0.038,
    'tax_rate': 0.22,
    'fringe_benefits_rate': 0.28,
    'personal_consumption_rate': 0.0,
    'post_injury_capacity': 'Limited to sedentary desk work only due to permanent lifting restrictions and chronic pain. Cannot perform previous engineering duties requiring site visits, equipment handling, or physical inspection activities.',
    'post_injury_occupation': 'Desk-based engineering consultant (limited availability)',
    'post_injury_annual_income': 40000.0,
    'labor_market_reduction': 0.65,
    'life_care_plan_cost': 150000.0,
    'household_services_cost': 95000.0
}

# Field descriptions and examples
FIELD_HELP = {
    'client_name': 'Full legal name of the injured party (e.g., "John Michael Smith")',
    'date_of_birth': 'Used to calculate age at injury and life expectancy tables',
    'date_of_injury': 'The specific date when the injury occurred - this starts the loss period',
    'date_of_report': 'Current report date - separates past vs future losses',
    'gender': 'Required for accurate life expectancy and work-life expectancy calculations',
    'education': 'Highest level completed (e.g., "High School", "Bachelor\'s in Engineering", "Master\'s in Business")',
    'marital_status': 'Current status (e.g., "Single", "Married", "Divorced", "Widowed")',
    'residence_state': 'State of primary residence for jurisdiction and wage data',
    'residence_county': 'County for local economic data and wage adjustments',
    'msa': 'Metropolitan Statistical Area for regional wage comparisons',
    'pre_injury_occupation': 'Detailed job title and duties (e.g., "Registered Nurse - ICU", "Construction Foreman")',
    'pre_injury_annual_income': 'Total annual earnings including overtime, bonuses, commissions before injury',
    'pre_injury_hourly_rate': 'Base hourly wage rate (calculated from annual if needed)',
    'employer': 'Company name and any relevant details about employment stability',
    'years_with_employer': 'Length of employment showing job stability and advancement potential',
    'injury_description': 'Detailed medical description: mechanism of injury, body parts affected, severity, initial treatment',
    'body_parts_injured': 'Specific anatomical areas (e.g., "Right shoulder, rotator cuff", "Lumbar spine L3-L5")',
    'medical_treatment': 'Comprehensive treatment history: surgeries, therapy, medications, ongoing care',
    'current_medical_status': 'Current condition, restrictions, prognosis, MMI status, permanent limitations',
    'life_expectancy': 'Total remaining years based on actuarial tables (from injury date)',
    'work_life_expectancy': 'Expected working years remaining (from injury date)',
    'years_to_final_separation': 'Additional non-working years (for work-life ratio calculations)',
    'wage_growth_rate': 'Annual percentage increase in wages (historical average 2.5-4%)',
    'discount_rate': 'Present value discount rate (typically government bond rates 3-5%)',
    'unemployment_rate': 'Regional unemployment percentage (reduces earning probability)',
    'tax_rate': 'Combined federal/state/local tax rate on earnings',
    'fringe_benefits_rate': 'Employer benefits as percentage of wages (health, retirement, etc.)',
    'personal_consumption_rate': 'Portion spent on self (used in wrongful death cases only)',
    'post_injury_capacity': 'Detailed description of remaining work abilities and limitations',
    'post_injury_occupation': 'Realistic alternative employment considering restrictions',
    'post_injury_annual_income': 'Projected annual earnings in alternative employment',
    'labor_market_reduction': 'Percentage reduction in earning capacity due to injury',
    'life_care_plan_cost': 'Total present value of future medical/care costs',
    'household_services_cost': 'Value of lost household services (cleaning, cooking, childcare)'
}

# Initialize session state for data persistence and calculations
if 'report_data' not in st.session_state:
    st.session_state.report_data = {}

if 'calculation_results' not in st.session_state:
    st.session_state.calculation_results = None

if 'show_results' not in st.session_state:
    st.session_state.show_results = False

# Load sample case function
def load_sample_case():
    """Load sample case data into session state"""
    st.session_state.report_data = SAMPLE_CASE.copy()
    st.session_state.show_results = False
    st.session_state.calculation_results = None
    st.success("Sample case loaded! Review the data in each tab and generate your report.")

@dataclass
class EconomicData:
    """Data structure for economic analysis"""
    # Personal Information
    client_name: str = ""
    date_of_birth: date = date(1980, 1, 1)
    date_of_injury: date = date(2023, 1, 1)
    date_of_report: date = date.today()
    age_at_injury: float = 0.0
    gender: str = "Male"
    education: str = ""
    marital_status: str = ""
    
    # Geographic Information
    residence_state: str = ""
    residence_county: str = ""
    msa: str = ""
    
    # Employment Information
    pre_injury_occupation: str = ""
    pre_injury_annual_income: float = 0.0
    pre_injury_hourly_rate: float = 0.0
    employer: str = ""
    years_with_employer: float = 0.0
    
    # Injury Information
    injury_description: str = ""
    body_parts_injured: str = ""
    medical_treatment: str = ""
    current_medical_status: str = ""
    
    # Work Life Expectancy Data
    life_expectancy: float = 78.5
    work_life_expectancy: float = 45.0
    years_to_final_separation: float = 0.0
    statistical_retirement_date: date = date.today()
    
    # Economic Factors
    wage_growth_rate: float = 0.03
    discount_rate: float = 0.04
    unemployment_rate: float = 0.035
    tax_rate: float = 0.12
    fringe_benefits_rate: float = 0.06
    personal_consumption_rate: float = 0.25
    
    # Post-Injury Information
    post_injury_capacity: str = ""
    post_injury_occupation: str = ""
    post_injury_annual_income: float = 0.0
    labor_market_reduction: float = 0.0
    
    # Additional Costs
    life_care_plan_cost: float = 0.0
    household_services_cost: float = 0.0

def calculate_age(birth_date: date, reference_date: date) -> float:
    """Calculate age in years with decimal precision"""
    return (reference_date - birth_date).days / 365.25

def calculate_statistical_dates(doi: date, le: float, wle: float) -> tuple:
    """Calculate statistical retirement and death dates"""
    retirement_date = doi + timedelta(days=int(wle * 365.25))
    death_date = doi + timedelta(days=int(le * 365.25))
    return retirement_date, death_date

def calculate_worklife_factor(wle: float, yfs: float) -> float:
    """Calculate worklife participation factor"""
    if yfs > 0:
        return wle / (wle + yfs)
    return 1.0

def calculate_adjusted_earnings_factor(unemployment: float, tax: float, 
                                     fringe: float = 0.0, consumption: float = 0.0,
                                     worklife_factor: float = 1.0) -> dict:
    """Calculate comprehensive adjusted earnings factor"""
    factors = []
    current_factor = 1.0
    
    # Worklife factor
    current_factor *= worklife_factor
    factors.append(("Worklife Factor", worklife_factor, current_factor))
    
    # Fringe benefits (additive)
    if fringe > 0:
        current_factor *= (1 + fringe)
        factors.append(("Fringe Benefits", 1 + fringe, current_factor))
    
    # Unemployment reduction
    current_factor *= (1 - unemployment)
    factors.append(("Unemployment Factor", 1 - unemployment, current_factor))
    
    # Tax reduction
    current_factor *= (1 - tax)
    factors.append(("Tax Factor", 1 - tax, current_factor))
    
    # Personal consumption (for wrongful death)
    if consumption > 0:
        current_factor = current_factor - consumption
        factors.append(("Personal Consumption", consumption, current_factor))
    
    return {
        'final_aef': current_factor,
        'factors': factors
    }

def generate_economic_schedule(data: EconomicData) -> tuple:
    """Generate past and future economic loss schedules"""
    
    # Calculate key dates and factors
    retirement_date, death_date = calculate_statistical_dates(
        data.date_of_injury, data.life_expectancy, data.work_life_expectancy
    )
    
    worklife_factor = calculate_worklife_factor(
        data.work_life_expectancy, data.years_to_final_separation
    )
    
    aef_result = calculate_adjusted_earnings_factor(
        data.unemployment_rate, data.tax_rate, 
        data.fringe_benefits_rate, data.personal_consumption_rate,
        worklife_factor
    )
    
    # Past losses (DOI to DOR)
    past_years = []
    report_year = data.date_of_report.year
    injury_year = data.date_of_injury.year
    
    for year in range(injury_year, report_year + 1):
        age = calculate_age(data.date_of_birth, date(year, 7, 1))
        years_since_injury = year - injury_year
        
        # Calculate portion of year
        if year == injury_year:
            # Partial year from injury date
            days_remaining = (date(year, 12, 31) - data.date_of_injury).days + 1
            portion = days_remaining / 365
        elif year == report_year:
            # Partial year to report date
            days_elapsed = data.date_of_report.timetuple().tm_yday
            portion = days_elapsed / 365
        else:
            portion = 1.0
        
        # Calculate earnings
        pre_injury_earnings = data.pre_injury_annual_income * ((1 + data.wage_growth_rate) ** years_since_injury)
        nominal_loss = pre_injury_earnings * portion
        adjusted_loss = nominal_loss * aef_result['final_aef']
        
        # Present value (discount back to injury date)
        pv_loss = adjusted_loss / ((1 + data.discount_rate) ** years_since_injury)
        
        past_years.append({
            'Year': year,
            'Age': round(age, 1),
            'Portion': round(portion * 100, 1),
            'Pre-Injury Earnings': round(pre_injury_earnings, 2),
            'Nominal Loss': round(nominal_loss, 2),
            'AEF': round(aef_result['final_aef'] * 100, 2),
            'Adjusted Loss': round(adjusted_loss, 2),
            'Present Value': round(pv_loss, 2)
        })
    
    # Future losses (DOR to retirement)
    future_years = []
    retirement_year = retirement_date.year
    
    for year in range(report_year, retirement_year + 1):
        if year == report_year:
            continue  # Already handled in past
            
        age = calculate_age(data.date_of_birth, date(year, 7, 1))
        years_since_injury = year - injury_year
        
        # Calculate portion of year
        if year == retirement_year:
            # Partial year to retirement
            retirement_day = retirement_date.timetuple().tm_yday
            portion = retirement_day / 365
        else:
            portion = 1.0
        
        # Calculate earnings
        pre_injury_earnings = data.pre_injury_annual_income * ((1 + data.wage_growth_rate) ** years_since_injury)
        post_injury_earnings = data.post_injury_annual_income * ((1 + data.wage_growth_rate) ** (years_since_injury - (report_year - injury_year)))
        
        nominal_loss = (pre_injury_earnings - post_injury_earnings) * portion
        adjusted_loss = nominal_loss * aef_result['final_aef']
        
        # Present value (discount to injury date)
        pv_loss = adjusted_loss / ((1 + data.discount_rate) ** years_since_injury)
        
        future_years.append({
            'Year': year,
            'Age': round(age, 1),
            'Portion': round(portion * 100, 1),
            'Pre-Injury Earnings': round(pre_injury_earnings, 2),
            'Post-Injury Earnings': round(post_injury_earnings, 2),
            'Nominal Loss': round(nominal_loss, 2),
            'AEF': round(aef_result['final_aef'] * 100, 2),
            'Adjusted Loss': round(adjusted_loss, 2),
            'Present Value': round(pv_loss, 2)
        })
    
    past_df = pd.DataFrame(past_years)
    future_df = pd.DataFrame(future_years)
    
    return past_df, future_df, aef_result, retirement_date, death_date

def create_word_report(data: EconomicData, past_df: pd.DataFrame, 
                      future_df: pd.DataFrame, aef_result: dict, 
                      retirement_date: date, death_date: date) -> bytes:
    """Generate comprehensive forensic economic expert report matching exact Kincaid Wolstein template"""
    
    doc = Document()
    
    # Calculate comprehensive totals
    past_total = past_df['Present Value'].sum() if not past_df.empty and 'Present Value' in past_df.columns else 0
    future_total = future_df['Present Value'].sum() if not future_df.empty and 'Present Value' in future_df.columns else 0
    total_economic_loss = past_total + future_total + data.life_care_plan_cost + data.household_services_cost
    
    # Set document margins
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    
    # TITLE PAGE - Exact Kincaid Wolstein format with professional logo
    header_table = doc.add_table(rows=1, cols=2)
    header_table.autofit = False
    header_table.columns[0].width = Inches(3.5)
    header_table.columns[1].width = Inches(4.5)
    
    # Left cell - Kincaid Wolstein logo styling (matching provided logo)
    left_cell = header_table.cell(0, 0)
    left_para = left_cell.paragraphs[0]
    left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Company name in exact blue color from logo
    company_run = left_para.add_run("Kincaid\nWolstein")
    company_run.font.name = 'Arial'
    company_run.font.size = Inches(0.35)
    company_run.font.color.rgb = RGBColor(0, 123, 193)  # Exact blue from logo
    company_run.bold = True
    
    # Add vertical line separator
    line_run = left_para.add_run(" | ")
    line_run.font.color.rgb = RGBColor(100, 100, 100)
    
    # Subtitle exactly as in logo
    subtitle_run = left_para.add_run("VOCATIONAL &\nREHABILITATION\nSERVICES")
    subtitle_run.font.name = 'Arial'
    subtitle_run.font.size = Inches(0.12)
    subtitle_run.font.color.rgb = RGBColor(100, 100, 100)
    
    # Right cell - Privilege notice
    right_cell = header_table.cell(0, 1)
    right_para = right_cell.paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    privilege_run = right_para.add_run("DRAFT-WORK PRODUCT PRIVILEGE")
    privilege_run.font.color.rgb = RGBColor(0, 123, 193)
    privilege_run.font.size = Inches(0.12)
    
    # Remove table borders
    for row in header_table.rows:
        for cell in row.cells:
            cell._element.get_or_add_tcPr().append(parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="nil"/><w:left w:val="nil"/><w:bottom w:val="nil"/><w:right w:val="nil"/></w:tcBorders>'))
    
    # Add spacing
    for _ in range(8):
        doc.add_paragraph()
    
    # Main title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("APPRAISAL OF ECONOMIC LOSS\nRESULTING FROM INJURY TO ")
    title_run.font.name = 'Arial'
    title_run.font.size = Inches(0.18)
    title_run.bold = True
    
    name_run = title_para.add_run(data.client_name.upper())
    name_run.font.name = 'Arial'
    name_run.font.size = Inches(0.18)
    name_run.bold = True
    name_run.font.color.rgb = RGBColor(0, 123, 193)
    
    # Add spacing
    for _ in range(3):
        doc.add_paragraph()
    
    # Prepared by section with exact formatting
    prepared_para = doc.add_paragraph()
    prepared_para.add_run("PREPARED BY:").bold = True
    prepared_para.add_run("\t\tKincaid Wolstein Vocational and Rehabilitation Services\n")
    prepared_para.add_run("\t\t\t\t\tOne University Plaza ~ Suite 302\n")
    prepared_para.add_run("\t\t\t\t\tHackensack, New Jersey 07601\n")
    prepared_para.add_run("\t\t\t\t\tPhone: (201) 343-0700\n")
    prepared_para.add_run("\t\t\t\t\tFax: (201) 343-0757")
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Case details
    prepared_for_para = doc.add_paragraph()
    prepared_for_para.add_run("PREPARED FOR:").bold = True
    
    doc.add_paragraph()
    
    regarding_para = doc.add_paragraph()
    regarding_para.add_run("REGARDING:").bold = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    dob_para = doc.add_paragraph()
    dob_para.add_run("DATE OF BIRTH:").bold = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    report_date_para = doc.add_paragraph()
    report_date_para.add_run("REPORT DATE:").bold = True
    
    doc.add_page_break()
    
    # PAGE 2 - TABLE OF CONTENTS
    add_page_header(doc, data.client_name, "2")
    
    toc_title = doc.add_heading('Table of Contents', level=1)
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Complete TOC from template
    toc_items = [
        "CERTIFICATION ................................................................................................................ 3",
        "PURPOSE OF APPRAISAL ............................................................................................... 4", 
        "OPINION OF ECONOMIC LOSSES.................................................................................. 4",
        "BACKGROUND FACTS AND ASSUMPTIONS ................................................................ 5",
        "Summary Information ..................................................................................................... 5",
        "Life Expectancy ............................................................................................................... 5",
        "Statistical retirement age ................................................................................................ 5",
        "Expected working years ................................................................................................. 5",
        "Worklife-to Retirement Ratio .......................................................................................... 6",
        "Description of Incident .................................................................................................... 6",
        "Occupation and Employment ......................................................................................... 6",
        "Earnings History .............................................................................................................. 7",
        "Pre-injury Earnings Capacity: ......................................................................................... 7",
        "Fringe Benefits ................................................................................................................ 7",
        "Functionality and Future Employability .......................................................................... 8",
        "Post-injury Earnings Capacity ........................................................................................ 9",
        "Growth Wage Rates ....................................................................................................... 9",
        "Household Services ...................................................................................................... 10",
        "COMPONENTS OF ANALYSIS ....................................................................................... 12",
        "PRE‐INJURY ADJUSTED EARNINGS ........................................................................... 13",
        "PRE‐INJURY ADJUSTED EARNINGS FOLLOWING RTW ........................................... 15",
        "POST‐INJURY ADJUSTED EARNINGS FOLLOWING DEPARTURE FROM FULL",
        "DUTY POSITION .............................................................................................................. 16",
        "COST OF LIFETIME CARE ............................................................................................. 18",
        "Methodology.................................................................................................................. 20",
        "Summary of Findings .................................................................................................... 20",
        "Important Considerations ............................................................................................. 21",
        "COST OF LIFETIME CARE ............................................................................................. 23",
        "TABLES .......................................................................................................................... 24",
        "STATEMENT OF ETHICAL PRINCIPLES AND PRINCIPLES OF PROFESSIONAL ... 38",
        "PRACTICE ...................................................................................................................... 38"
    ]
    
    for item in toc_items:
        doc.add_paragraph(item)
    
    add_page_footer(doc)
    doc.add_page_break()
    
    # PAGE 3 - ECONOMIC LOSS APPRAISAL REPORT
    add_page_header(doc, data.client_name, "3")
    
    doc.add_heading('ECONOMIC LOSS APPRAISAL REPORT', level=1)
    doc.add_heading('Introduction', level=2)
    
    intro_text = f"""This Economic Loss Appraisal Report has been prepared by [Expert Name], a forensic economist, in the matter of [Case Name], to evaluate the economic losses resulting from {data.client_name}'s injuries (or wrongful death). The purpose of this report is to quantify various categories of economic damages in a manner that is clear, comprehensive, and compliant with the Daubert standard for expert testimony. The key components of loss addressed include:

• Lost Wages/Earnings Capacity – the value of past and future income (and benefits) lost due to the incident.
• Future Medical and Healthcare Costs – the present value of reasonable future medical expenses related to the injury.
• Loss of Household Services – the economic value of household tasks and services {data.client_name} can no longer perform.

All findings are presented in plain language with technical details explained for a general audience. Specialized terms are defined throughout the report and in a glossary. The methodologies used are grounded in well-established economic principles and reliable methods that have been published in peer-reviewed literature and are generally accepted in the field. This report adheres to Daubert criteria by using reliable principles and methods, referencing peer-reviewed sources, discussing potential error rates or uncertainties, and applying the methods to the facts of this case. All assumptions, data sources, and calculations are documented in the sections below, and all opinions are stated to a reasonable degree of economic certainty.

Report Structure: After summarizing the data and assumptions considered, the report details each category of loss in turn (Lost Wages, Future Medical Costs, Household Services), followed by a summary of total losses. Each section outlines the methodology and reasoning, ensuring transparency and layperson accessibility. A final section provides requisite expert disclosures, including qualifications, materials reviewed, assumptions, exhibits, and a signature attestation."""
    
    doc.add_paragraph(intro_text)
    
    # Materials Considered
    doc.add_heading('Materials Considered', level=2)
    materials_text = f"""In preparing this analysis, I have reviewed and relied upon the following materials and data (among others):

• Case Documents: [List key documents: e.g. Complaint, deposition of {data.client_name}, accident reports].
• Medical Records & Life Care Plan: [e.g. Medical reports from Dr. A; Life Care Plan by [Name] dated ___ outlining future care needs].
• Employment & Earnings Records: [e.g. pay stubs, W-2 forms, tax returns, employment file from {data.client_name}'s employer].
• Economic & Statistical Data: U.S. Bureau of Labor Statistics (wage data, Consumer Price Index), published worklife expectancy tables, life expectancy tables (e.g. U.S. CDC Life Tables), and relevant economic research literature.
• Other: [Any other data: e.g. vocational expert report, family testimony on household services, etc.].

(Modify the above list as needed for the specific case, ensuring all materials considered are listed.)"""
    
    doc.add_paragraph(materials_text)
    
    # Key Assumptions with dynamic data
    doc.add_heading('Key Assumptions', level=2)
    assumptions_text = f"""The following assumptions underlie the calculations in this report (to be adjusted per case specifics):

• Employment but for Incident: It is assumed that absent the incident, {data.client_name} would have continued working in their usual capacity up to a normal retirement age of {calculate_age(data.date_of_birth, retirement_date):.0f} years (or for {data.work_life_expectancy:.1f} additional years of worklife). This worklife expectancy is based on statistical averages adjusted for {data.client_name}'s age, gender, and work history.

• Post-Incident Work Capacity: {data.client_name} is now unable to work (or can only work in a limited capacity earning ${data.post_injury_annual_income:,.2f} per year), per medical/vocational evidence, resulting in a loss of earning capacity as detailed below.

• Fringe Benefits: Employer-provided benefits (health insurance, retirement contributions, etc.) comprised approximately {data.fringe_benefits_rate*100:.1f}% of wages and are included as part of the lost compensation.

• Income Taxes: Calculations of lost earnings are presented on an [after-tax basis] (if applicable under jurisdiction) so that the award reflects net take-home pay loss. A combined federal/state tax rate of {data.tax_rate*100:.1f}% is assumed for this purpose.

• Life Expectancy: {data.client_name} has a remaining life expectancy of {data.life_expectancy:.1f} years, based on [source, e.g. U.S. Life Tables], which is used to project future losses through the year {death_date.year}. If the injury is expected to impact longevity, medical expert input is used to adjust this expectation.

• Discount Rate: A discount rate of {data.discount_rate*100:.1f}% is used to convert future dollars to present value. This rate is chosen to reflect a risk-free rate of return (e.g. based on U.S. Treasury or tax-free municipal bond yields) appropriate for a lump-sum award.

• Inflation Rates: Future wage growth and medical cost inflation are assumed at {data.wage_growth_rate*100:.1f}% annually (based on historical data and current economic forecasts for wages and healthcare costs, respectively). These inflation assumptions are paired with the discount rate to ensure consistency (e.g. real vs. nominal projections).

• Household Services: It is assumed that {data.client_name} performed approximately ___ hours per week of household and family services pre-incident (based on time-use data or family testimony), and that due to the injury [he/she] can no longer perform [all or a specific portion] of these tasks. The types of services affected include [list: e.g. cleaning, cooking, childcare, yard work].

• Mitigation: Any mitigation or offset (such as actual earnings post-incident, or replacement services provided by others) has been considered and will be noted in the calculations. It is assumed that {data.client_name} has made reasonable efforts to mitigate losses where possible.

(Each assumption should be reviewed and modified to fit the facts of the case. Additional assumptions can be added as needed.)"""
    
    doc.add_paragraph(assumptions_text)
    
    add_page_footer(doc)
    doc.add_page_break()
    
    # Continue implementing all remaining pages...
    # For brevity, adding key sections
    
    # Final page - Save document
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io.getvalue()

def add_page_header(doc, client_name, page_num):
    """Add consistent page header with client name and page number"""
    header_para = doc.add_paragraph()
    names = client_name.split()
    if len(names) >= 2:
        header_text = f"{names[-1].upper()}, {names[0].upper()}"
    else:
        header_text = client_name.upper()
    
    header_para.add_run(header_text)
    header_para.add_run("\t\t\t")
    header_para.add_run("DRAFT-WORK PRODUCT PRIVILEGE")
    header_para.add_run(f"\t\t\t\t\t\t\t\t{page_num}")
    
def add_page_footer(doc):
    """Add consistent page footer with company information"""
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.add_run("Kincaid Wolstein Vocational and Rehabilitation Services\n")
    footer_para.add_run("1 University Plaza, Suite 302, Hackensack, New Jersey 07601\n")
    footer_para.add_run("www.KWVRS.com ■ Info@KWVRS.Com ■ Tel: (201) 343-0700")

def create_word_report_original(data: EconomicData, past_df: pd.DataFrame, 
                      future_df: pd.DataFrame, aef_result: dict, 
                      retirement_date: date, death_date: date) -> bytes:
    """Generate comprehensive Word report following the complete template structure"""
    
    doc = Document()
    
    # Calculate totals for use throughout the document
    past_total = past_df['Present Value'].sum() if not past_df.empty and 'Present Value' in past_df.columns else 0
    future_total = future_df['Present Value'].sum() if not future_df.empty and 'Present Value' in future_df.columns else 0
    
    # Set document margins and orientation
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    
    # Title page header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header.add_run("DRAFT-WORK PRODUCT PRIVILEGE")
    header_run.bold = True
    
    # Add spacing
    for _ in range(8):
        doc.add_paragraph()
    
    # Main title
    title = doc.add_heading('APPRAISAL OF ECONOMIC LOSS', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading(f'RESULTING FROM INJURY TO {data.client_name.upper()}', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add spacing
    for _ in range(3):
        doc.add_paragraph()
    
    # Prepared by section
    prepared_by = doc.add_paragraph()
    prepared_by.alignment = WD_ALIGN_PARAGRAPH.LEFT
    prepared_by.add_run("PREPARED BY:\n").bold = True
    prepared_by.add_run("Kincaid Wolstein Vocational and Rehabilitation Services\n")
    prepared_by.add_run("One University Plaza ~ Suite 302\n")
    prepared_by.add_run("Hackensack, New Jersey 07601\n")
    prepared_by.add_run("Phone: (201) 343-0700\n")
    prepared_by.add_run("Fax: (201) 343-0757\n")
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Case details
    prepared_for = doc.add_paragraph()
    prepared_for.add_run("PREPARED FOR:\n").bold = True
    prepared_for.add_run("[Attorney Name/Firm]\n")
    
    doc.add_paragraph()
    
    regarding = doc.add_paragraph()
    regarding.add_run("REGARDING:\n").bold = True
    regarding.add_run("Personal Injury Economic Loss Analysis\n")
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    dob = doc.add_paragraph()
    dob.add_run("DATE OF BIRTH:\n").bold = True
    dob.add_run(f"{data.date_of_birth.strftime('%B %d, %Y')}\n")
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    report_date = doc.add_paragraph()
    report_date.add_run("REPORT DATE:\n").bold = True
    report_date.add_run(f"{data.date_of_report.strftime('%B %d, %Y')}\n")
    
    doc.add_paragraph()
    
    # Case details
    case_details = [
        ("PREPARED FOR:", ""),
        ("REGARDING:", data.client_name),
        ("DATE OF BIRTH:", data.date_of_birth.strftime('%m/%d/%Y')),
        ("REPORT DATE:", data.date_of_report.strftime('%m/%d/%Y'))
    ]
    
    for label, value in case_details:
        para = doc.add_paragraph()
        para.add_run(label).bold = True
        para.add_run(f"\t\t{value}")
    
    doc.add_page_break()
    
    # Table of Contents
    toc_heading = doc.add_heading('Table of Contents', level=1)
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    toc_items = [
        "CERTIFICATION\t3",
        "PURPOSE OF APPRAISAL\t4", 
        "OPINION OF ECONOMIC LOSSES\t4",
        "BACKGROUND FACTS AND ASSUMPTIONS\t5",
        "Summary Information\t5",
        "Life Expectancy\t5",
        "Statistical retirement age\t5",
        "Expected working years\t5",
        "Worklife-to Retirement Ratio\t6",
        "Description of Incident\t6",
        "Occupation and Employment\t6",
        "Earnings History\t7",
        "Pre-injury Earnings Capacity\t7",
        "Fringe Benefits\t7",
        "Functionality and Future Employability\t8",
        "Post-injury Earnings Capacity\t9",
        "Growth Wage Rates\t9",
        "Household Services\t10",
        "COMPONENTS OF ANALYSIS\t12",
        "PRE‐INJURY ADJUSTED EARNINGS\t13",
        "PRE‐INJURY ADJUSTED EARNINGS FOLLOWING RTW\t15",
        "POST‐INJURY ADJUSTED EARNINGS FOLLOWING DEPARTURE FROM FULL DUTY POSITION\t16",
        "COST OF LIFETIME CARE\t18",
        "Methodology\t20",
        "Summary of Findings\t20",
        "Important Considerations\t21",
        "COST OF LIFETIME CARE\t23",
        "TABLES\t24",
        "STATEMENT OF ETHICAL PRINCIPLES AND PRINCIPLES OF PROFESSIONAL PRACTICE\t38"
    ]
    
    for item in toc_items:
        doc.add_paragraph(item)
    
    doc.add_page_break()
    
    # CERTIFICATION
    doc.add_heading('CERTIFICATION', level=1)
    cert_text = f"""
I certify that I have prepared this economic analysis in accordance with the principles and standards of practice in forensic economics. The opinions expressed in this report are based on my professional expertise and are offered to a reasonable degree of economic certainty.

The methodologies employed are consistent with those generally accepted in the field of forensic economics and have been subjected to peer review in professional literature. All assumptions, data sources, and calculations are documented and transparent.

This report contains my complete professional opinions on the economic losses sustained by {data.client_name} as a result of the injuries sustained on {data.date_of_injury.strftime('%B %d, %Y')}.

_________________________________
Forensic Economist
Date: {data.date_of_report.strftime('%B %d, %Y')}
"""
    doc.add_paragraph(cert_text)
    
    doc.add_page_break()
    
    # PURPOSE OF APPRAISAL
    doc.add_heading('PURPOSE OF APPRAISAL', level=1)
    purpose_text = f"""
This Economic Loss Appraisal Report has been prepared to evaluate the economic losses resulting from {data.client_name}'s injuries sustained on {data.date_of_injury.strftime('%B %d, %Y')}. The purpose of this report is to quantify various categories of economic damages in a manner that is clear, comprehensive, and compliant with the Daubert standard for expert testimony.

The key components of loss addressed include:
• Lost Wages/Earnings Capacity – the value of past and future income (and benefits) lost due to the incident
• Future Medical and Healthcare Costs – the present value of reasonable future medical expenses related to the injury
• Loss of Household Services – the economic value of household tasks and services {data.client_name} can no longer perform

All findings are presented in plain language with technical details explained for a general audience. The methodologies used are grounded in well-established economic principles and reliable methods that have been published in peer-reviewed literature and are generally accepted in the field.
"""
    doc.add_paragraph(purpose_text)
    
    # OPINION OF ECONOMIC LOSSES
    doc.add_heading('OPINION OF ECONOMIC LOSSES', level=1)
    
    past_total = past_df['Present Value'].sum() if not past_df.empty else 0
    future_total = future_df['Present Value'].sum() if not future_df.empty else 0
    total_loss = past_total + future_total + data.life_care_plan_cost + data.household_services_cost
    
    opinion_text = f"""
Based on my analysis of the economic impact of {data.client_name}'s injuries, the total present value of economic losses is:

TOTAL ECONOMIC LOSS: ${total_loss:,.2f}

This total comprises the following components:
• Past Lost Earnings (Injury to Report Date): ${past_total:,.2f}
• Future Lost Earnings (Report Date to Retirement): ${future_total:,.2f}
• Life Care Plan Costs: ${data.life_care_plan_cost:,.2f}
• Household Services Loss: ${data.household_services_cost:,.2f}

These calculations employ established economic methodologies including present value analysis, demographic adjustments, and comprehensive adjustment factors. The analysis incorporates the Tinari Adjustment methodology with a final Adjusted Earnings Factor of {aef_result['final_aef']:.4f}.

All amounts are expressed in present value as of {data.date_of_report.strftime('%B %d, %Y')}, representing the lump sum that would be required today to compensate for the identified economic losses.
"""
    doc.add_paragraph(opinion_text)
    
    doc.add_page_break()
    
    # BACKGROUND FACTS AND ASSUMPTIONS
    doc.add_heading('BACKGROUND FACTS AND ASSUMPTIONS', level=1)
    
    # Summary Information
    doc.add_heading('Summary Information', level=2)
    summary_table = doc.add_table(rows=15, cols=2)
    summary_table.style = 'Table Grid'
    
    summary_info_data = [
        ('Client Name', data.client_name),
        ('Date of Birth', data.date_of_birth.strftime('%m/%d/%Y')),
        ('Date of Injury', data.date_of_injury.strftime('%m/%d/%Y')),
        ('Date of Report', data.date_of_report.strftime('%m/%d/%Y')),
        ('Age at Injury', f"{calculate_age(data.date_of_birth, data.date_of_injury):.1f} years"),
        ('Current Age', f"{calculate_age(data.date_of_birth, data.date_of_report):.1f} years"),
        ('Gender', data.gender),
        ('Education', data.education),
        ('Marital Status', data.marital_status),
        ('Residence', f"{data.residence_county}, {data.residence_state}"),
        ('MSA', data.msa),
        ('Pre-Injury Occupation', data.pre_injury_occupation),
        ('Employer', data.employer),
        ('Years with Employer', f"{data.years_with_employer:.1f} years"),
        ('Pre-Injury Annual Income', f"${data.pre_injury_annual_income:,.2f}")
    ]
    
    for i, (label, value) in enumerate(summary_info_data):
        summary_table.cell(i, 0).text = label
        summary_table.cell(i, 1).text = str(value)
    
    # Life Expectancy
    doc.add_heading('Life Expectancy', level=2)
    le_text = f"""
Based on standard actuarial tables and demographic data:
• Life Expectancy from Date of Injury: {data.life_expectancy:.1f} years
• Statistical Death Date: {death_date.strftime('%B %d, %Y')}
• Remaining Life Expectancy from Report Date: {(death_date - data.date_of_report).days / 365.25:.1f} years

These projections are based on U.S. Life Tables published by the Centers for Disease Control and Prevention, adjusted for the individual's age, gender, and other demographic characteristics.
"""
    doc.add_paragraph(le_text)
    
    # Statistical retirement age
    doc.add_heading('Statistical Retirement Age', level=2)
    retirement_text = f"""
Work-life expectancy analysis based on labor force participation data:
• Work-Life Expectancy from Date of Injury: {data.work_life_expectancy:.1f} years
• Statistical Retirement Date: {retirement_date.strftime('%B %d, %Y')}
• Remaining Work Years from Report Date: {(retirement_date - data.date_of_report).days / 365.25:.1f} years

Work-life expectancy projections are derived from U.S. Bureau of Labor Statistics data on labor force participation rates, accounting for age, gender, education, and occupation-specific patterns.
"""
    doc.add_paragraph(retirement_text)
    
    # Expected working years
    doc.add_heading('Expected Working Years', level=2)
    total_work_years = (retirement_date - data.date_of_injury).days / 365.25
    past_work_years = (data.date_of_report - data.date_of_injury).days / 365.25
    future_work_years = (retirement_date - data.date_of_report).days / 365.25
    
    work_years_text = f"""
Total Economic Analysis Period: {total_work_years:.1f} years
• Past Period (Injury to Report Date): {past_work_years:.1f} years
• Future Period (Report Date to Retirement): {future_work_years:.1f} years

This analysis covers the entire period from the date of injury through the projected statistical retirement date, encompassing both documented past losses and projected future losses.
"""
    doc.add_paragraph(work_years_text)
    
    # Worklife-to Retirement Ratio
    doc.add_heading('Worklife-to Retirement Ratio', level=2)
    worklife_factor = calculate_worklife_factor(data.work_life_expectancy, data.years_to_final_separation)
    ratio_text = f"""
Worklife Participation Factor: {worklife_factor:.4f}

This factor represents the probability of continued labor force participation through the projected work-life period. It accounts for typical patterns of labor force entry and exit, including voluntary and involuntary separations from employment.

Calculation: Work-Life Expectancy ÷ (Work-Life Expectancy + Years to Final Separation)
{data.work_life_expectancy:.1f} ÷ ({data.work_life_expectancy:.1f} + {data.years_to_final_separation:.1f}) = {worklife_factor:.4f}
"""
    doc.add_paragraph(ratio_text)
    
    # Description of Incident
    doc.add_heading('Description of Incident', level=2)
    incident_text = f"""
Date of Injury: {data.date_of_injury.strftime('%B %d, %Y')}

Nature of Injury:
{data.injury_description}

Body Parts Injured:
{data.body_parts_injured}

Medical Treatment Received:
{data.medical_treatment}

Current Medical Status:
{data.current_medical_status}

The injury has resulted in significant functional limitations that impact {data.client_name}'s ability to perform the essential functions of their pre-injury occupation and has reduced their overall earning capacity in the labor market.
"""
    doc.add_paragraph(incident_text)
    
    # Occupation and Employment
    doc.add_heading('Occupation and Employment', level=2)
    employment_text = f"""
PRE-INJURY EMPLOYMENT:
• Occupation: {data.pre_injury_occupation}
• Employer: {data.employer}
• Length of Employment: {data.years_with_employer:.1f} years
• Annual Income: ${data.pre_injury_annual_income:,.2f}
• Hourly Rate: ${data.pre_injury_hourly_rate:.2f}

Employment History and Stability:
The employment record demonstrates stable, long-term employment with consistent earning capacity. The position required physical capabilities that are no longer available due to the injury.

POST-INJURY WORK CAPACITY:
{data.post_injury_capacity}

POST-INJURY EMPLOYMENT POTENTIAL:
• Limited Occupation: {data.post_injury_occupation}
• Projected Annual Income: ${data.post_injury_annual_income:,.2f}
• Labor Market Reduction: {data.labor_market_reduction*100:.1f}%

The significant reduction in earning capacity reflects both the functional limitations imposed by the injury and the competitive disadvantage in the labor market for available alternative employment.
"""
    doc.add_paragraph(employment_text)
    
    # Earnings History
    doc.add_heading('Earnings History', level=2)
    earnings_history_text = f"""
Pre-injury earnings analysis is based on documented income from employment records:

• Base Annual Earnings: ${data.pre_injury_annual_income:,.2f}
• Hourly Rate: ${data.pre_injury_hourly_rate:.2f}
• Employment Duration: {data.years_with_employer:.1f} years

Earnings demonstrate a stable pattern of income generation with potential for continued growth based on historical wage patterns and career progression typical for the occupation and industry.
"""
    doc.add_paragraph(earnings_history_text)
    
    # Pre-injury Earnings Capacity
    doc.add_heading('Pre-injury Earnings Capacity', level=2)
    pre_capacity_text = f"""
The pre-injury earnings capacity represents the income {data.client_name} would have earned absent the injury, including:

• Base annual income of ${data.pre_injury_annual_income:,.2f}
• Projected annual wage growth of {data.wage_growth_rate*100:.1f}%
• Continued employment through statistical retirement at age {calculate_age(data.date_of_birth, retirement_date):.0f}

This earning capacity is projected through the expected work-life period using established economic growth assumptions and demographic data on labor force participation.
"""
    doc.add_paragraph(pre_capacity_text)
    
    # Fringe Benefits
    doc.add_heading('Fringe Benefits', level=2)
    fringe_text = f"""
Employer-provided fringe benefits comprise {data.fringe_benefits_rate*100:.1f}% of total compensation and include:

• Health insurance coverage
• Retirement plan contributions
• Life and disability insurance
• Paid time off and holidays
• Other employer-provided benefits

The value of these benefits is included in the total compensation analysis as they represent real economic value that would have been received but for the injury. The fringe benefit rate is based on U.S. Bureau of Labor Statistics data on employee compensation.
"""
    doc.add_paragraph(fringe_text)
    
    # Functionality and Future Employability
    doc.add_heading('Functionality and Future Employability', level=2)
    functionality_text = f"""
Post-Injury Functional Capacity:
{data.post_injury_capacity}

Impact on Employability:
The functional limitations resulting from the injury significantly impact {data.client_name}'s ability to:
• Perform the essential functions of their pre-injury occupation
• Compete effectively in the labor market
• Maintain the same level of earning capacity

Vocational analysis indicates a substantial reduction in labor market access and earning potential due to the physical and functional limitations imposed by the injury.
"""
    doc.add_paragraph(functionality_text)
    
    # Post-injury Earnings Capacity
    doc.add_heading('Post-injury Earnings Capacity', level=2)
    post_capacity_text = f"""
Given the functional limitations, {data.client_name}'s post-injury earning capacity is significantly reduced:

• Alternative occupation: {data.post_injury_occupation}
• Projected annual income: ${data.post_injury_annual_income:,.2f}
• Reduction from pre-injury income: ${data.pre_injury_annual_income - data.post_injury_annual_income:,.2f} ({((data.pre_injury_annual_income - data.post_injury_annual_income) / data.pre_injury_annual_income * 100):.1f}%)

This represents the residual earning capacity available in the labor market considering the injury-related limitations and competitive disadvantages.
"""
    doc.add_paragraph(post_capacity_text)
    
    # Growth Wage Rates
    doc.add_heading('Growth Wage Rates', level=2)
    growth_text = f"""
Wage Growth Assumptions:
• Annual wage growth rate: {data.wage_growth_rate*100:.1f}%
• Basis: Historical wage growth data and economic forecasting
• Application: Applied to both pre-injury and post-injury earning capacity

The wage growth rate is based on long-term historical averages from the U.S. Bureau of Labor Statistics and reflects typical productivity and inflation-related wage increases over time. This rate is applied consistently to project future earning capacity through the work-life period.
"""
    doc.add_paragraph(growth_text)
    
    # Household Services
    if data.household_services_cost > 0:
        doc.add_heading('Household Services', level=2)
        household_text = f"""
Loss of Household Services Analysis:

The injury has resulted in a diminished capacity to perform household and family services, representing a real economic loss that must be quantified using replacement cost methodology.

• Total Present Value of Lost Household Services: ${data.household_services_cost:,.2f}

This analysis considers:
• Types of services no longer performable due to injury limitations
• Frequency and duration of service requirements
• Market rates for replacement services
• Present value calculation over life expectancy

The valuation uses the replacement cost method, determining the market cost to hire service providers to perform the household tasks that can no longer be completed due to the injury.
"""
        doc.add_paragraph(household_text)
    
    doc.add_page_break()
    
    # COMPONENTS OF ANALYSIS
    doc.add_heading('COMPONENTS OF ANALYSIS', level=1)
    
    components_text = f"""
This economic analysis consists of several key components that together quantify the total economic impact of {data.client_name}'s injuries:

1. PAST ECONOMIC LOSSES
   Analysis of lost earnings from the date of injury ({data.date_of_injury.strftime('%B %d, %Y')}) through the date of report ({data.date_of_report.strftime('%B %d, %Y')})

2. FUTURE ECONOMIC LOSSES
   Projection of lost earning capacity from the report date through statistical retirement ({retirement_date.strftime('%B %d, %Y')})

3. ADJUSTED EARNINGS FACTOR (AEF)
   Comprehensive adjustment methodology incorporating:
   • Worklife participation factors
   • Unemployment contingencies
   • Tax considerations
   • Fringe benefit inclusions
   • Personal consumption adjustments (if applicable)

4. PRESENT VALUE ANALYSIS
   Conversion of all future losses to present value using appropriate discount rates

5. ADDITIONAL ECONOMIC IMPACTS
   • Life care plan costs
   • Household services losses
   • Other quantifiable economic impacts

Each component employs established methodologies consistent with accepted practices in forensic economics.
"""
    doc.add_paragraph(components_text)
    
    doc.add_page_break()
    
    # PRE-INJURY ADJUSTED EARNINGS
    doc.add_heading('PRE‐INJURY ADJUSTED EARNINGS', level=1)
    
    pre_earnings_text = f"""
Pre-injury adjusted earnings represent the income {data.client_name} would have earned but for the injury, adjusted for real-world economic factors:

BASE EARNINGS CALCULATION:
• Annual pre-injury income: ${data.pre_injury_annual_income:,.2f}
• Projected annual growth: {data.wage_growth_rate*100:.1f}%
• Work-life period: {(retirement_date - data.date_of_injury).days / 365.25:.1f} years

ADJUSTMENT FACTOR METHODOLOGY:
The Adjusted Earnings Factor (AEF) of {aef_result['final_aef']:.4f} is applied to account for:
"""
    
    for factor_name, factor_value, cumulative in aef_result['factors']:
        pre_earnings_text += f"\n• {factor_name}: {factor_value:.4f}"
    
    pre_earnings_text += f"""

FINAL ADJUSTED EARNINGS:
Gross earnings are multiplied by the AEF to determine the net economic impact. This methodology ensures that the analysis reflects realistic earning patterns and economic conditions rather than theoretical maximum earnings.

The pre-injury adjusted earnings form the baseline for calculating economic losses throughout the analysis period.
"""
    doc.add_paragraph(pre_earnings_text)
    
    # Add the detailed past losses table
    if not past_df.empty:
        doc.add_page_break()
        doc.add_heading('PAST ECONOMIC LOSSES - DETAILED ANALYSIS', level=1)
        
        past_analysis_text = f"""
Detailed analysis of past economic losses from {data.date_of_injury.strftime('%B %d, %Y')} to {data.date_of_report.strftime('%B %d, %Y')}:

ANALYSIS METHODOLOGY:
• Year-by-year calculation of lost earning capacity
• Application of AEF adjustments
• Present value discounting to report date
• Proportional calculation for partial years

TOTAL PAST LOSSES: ${past_total:,.2f}
"""
        doc.add_paragraph(past_analysis_text)
        
        if not past_df.empty:
            # Create past losses table
            past_table = doc.add_table(rows=len(past_df) + 2, cols=len(past_df.columns))
            past_table.style = 'Table Grid'
            
            # Headers
            for j, col in enumerate(past_df.columns):
                cell = past_table.cell(0, j)
                cell.text = col
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # Data
            for i in range(len(past_df)):
                for j, col in enumerate(past_df.columns):
                    value = past_df.iloc[i, j]
                    if isinstance(value, (int, float)) and col in ['Pre-Injury Earnings', 'Nominal Loss', 'Adjusted Loss', 'Present Value']:
                        past_table.cell(i + 1, j).text = f"${value:,.2f}"
                    else:
                        past_table.cell(i + 1, j).text = str(value)
            
            # Totals row
            totals_row = len(past_df) + 1
            past_table.cell(totals_row, 0).text = "TOTALS"
            for j, col in enumerate(past_df.columns):
                if col in ['Pre-Injury Earnings', 'Nominal Loss', 'Adjusted Loss', 'Present Value']:
                    total = past_df[col].sum()
                    cell = past_table.cell(totals_row, j)
                    cell.text = f"${total:,.2f}"
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
    
    # Add future losses analysis
    if not future_df.empty:
        doc.add_page_break()
        doc.add_heading('FUTURE ECONOMIC LOSSES - DETAILED PROJECTIONS', level=1)
        
        future_analysis_text = f"""
Projected future economic losses from {data.date_of_report.strftime('%B %d, %Y')} to {retirement_date.strftime('%B %d, %Y')}:

PROJECTION METHODOLOGY:
• Pre-injury earning capacity with growth projections
• Post-injury residual earning capacity
• Net loss calculation (pre-injury minus post-injury)
• AEF adjustments applied consistently
• Present value discounting to report date

KEY ASSUMPTIONS:
• Wage growth rate: {data.wage_growth_rate*100:.1f}% annually
• Discount rate: {data.discount_rate*100:.1f}% annually
• Work-life period: {(retirement_date - data.date_of_report).days / 365.25:.1f} years

TOTAL FUTURE LOSSES: ${future_total:,.2f}
"""
        doc.add_paragraph(future_analysis_text)
        
        if not future_df.empty:
            # Create future losses table
            future_table = doc.add_table(rows=len(future_df) + 2, cols=len(future_df.columns))
            future_table.style = 'Table Grid'
            
            # Headers
            for j, col in enumerate(future_df.columns):
                cell = future_table.cell(0, j)
                cell.text = col
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # Data
            for i in range(len(future_df)):
                for j, col in enumerate(future_df.columns):
                    value = future_df.iloc[i, j]
                    if isinstance(value, (int, float)) and col in ['Pre-Injury Earnings', 'Post-Injury Earnings', 'Nominal Loss', 'Adjusted Loss', 'Present Value']:
                        future_table.cell(i + 1, j).text = f"${value:,.2f}"
                    else:
                        future_table.cell(i + 1, j).text = str(value)
            
            # Totals row
            totals_row = len(future_df) + 1
            future_table.cell(totals_row, 0).text = "TOTALS"
            for j, col in enumerate(future_df.columns):
                if col in ['Pre-Injury Earnings', 'Post-Injury Earnings', 'Nominal Loss', 'Adjusted Loss', 'Present Value']:
                    total = future_df[col].sum()
                    cell = future_table.cell(totals_row, j)
                    cell.text = f"${total:,.2f}"
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
    
    # Life Care Plan section
    if data.life_care_plan_cost > 0:
        doc.add_page_break()
        doc.add_heading('COST OF LIFETIME CARE', level=1)
        
        lcp_text = f"""
Life Care Plan Cost Analysis:

The life care plan represents the present value of future medical and care-related expenses resulting from the injury. This analysis is based on medical expert evaluation of future care needs and associated costs.

METHODOLOGY:
• Medical necessity assessment by qualified professionals
• Current cost documentation for required services and equipment
• Medical inflation projections for future cost growth
• Present value calculation using appropriate discount rates
• Life expectancy considerations for duration of care

TOTAL LIFE CARE PLAN COST (Present Value): ${data.life_care_plan_cost:,.2f}

This amount represents the funds needed today to cover all projected future medical expenses, therapies, equipment, and care services related to the injury over {data.client_name}'s remaining lifetime.

COST CATEGORIES:
The life care plan typically includes:
• Future medical treatments and surgeries
• Ongoing therapy and rehabilitation services
• Medical equipment and assistive devices
• Prescription medications
• Home health and personal care services
• Transportation for medical appointments
• Home and vehicle modifications

All costs are projected using appropriate medical inflation rates and discounted to present value to ensure adequate funding for future care needs.
"""
        doc.add_paragraph(lcp_text)
    
    # Methodology section
    doc.add_page_break()
    doc.add_heading('METHODOLOGY', level=1)
    
    methodology_text = f"""
This economic analysis employs established methodologies that meet standards for scientific reliability and admissibility under Daubert criteria:

1. PRESENT VALUE CONCEPT
Present Value calculations account for two key factors: (1) expected inflation (the rising cost of wages or medical services over time), and (2) the time value of money (the interest or return that can be earned by investing the award). Present value calculations essentially convert future dollar amounts to today's dollars using an appropriate discount rate.

2. RELIABLE PRINCIPLES & METHODS
The methods used in this analysis are standard in the field of forensic economics and based on well-established financial principles. These methods have been published in economic literature and subjected to peer review. The approach to valuing lost earnings and benefits draws on published models, including Dr. Frank Tinari's algebraic method for lost earnings (Tinari, F.D. 2016. "Demonstrating Lost Earnings: Algebraic vs. Spreadsheet Method." The Earnings Analyst, 15, 21–32).

3. KNOWN ERROR RATES & UNCERTAINTY
Economic damage assessment is based on forecasts and statistical averages; as such, it is not amenable to a precise "error rate" in the same sense as a laboratory experiment. However, to address uncertainty, conservative assumptions have been used and sensitivity analyses can be performed if needed. The data inputs (wage rates, growth rates) are derived from large samples or reliable studies, reducing the likelihood of significant error.

4. APPLICATION TO CASE FACTS
The principles and models have been applied specifically to the facts of {data.client_name}'s situation. All calculations use case-specific data (such as {data.client_name}'s actual earnings, medical needs, etc.), and the assumptions are tailored to {data.client_name} (taking into account age, health status, occupation, etc.).

5. DATA SOURCES
• Economic & Statistical Data: U.S. Bureau of Labor Statistics (wage data, Consumer Price Index)
• Published worklife expectancy tables and life expectancy tables (U.S. CDC Life Tables)
• Relevant economic research literature
• Standard actuarial and demographic data sources

6. ADJUSTMENT FACTOR METHODOLOGY
The analysis employs Frank Tinari's algebraic methodology for demonstrating lost earnings, which condenses calculations into a clear formula format. The Adjusted Earnings Factor (AEF) methodology incorporates:
   • Worklife participation factors
   • Unemployment contingencies ({data.unemployment_rate*100:.1f}%)
   • Tax considerations ({data.tax_rate*100:.1f}%)
   • Fringe benefit inclusions ({data.fringe_benefits_rate*100:.1f}%)
   • Personal consumption adjustments (if applicable)

7. PROFESSIONAL STANDARDS
This analysis follows accepted practices in forensic economics and methodology consistent with standards for expert economic testimony. All assumptions and data sources are documented and transparent, and results comply with professional ethical standards for economic analysis.
"""
    doc.add_paragraph(methodology_text)
    
    # Summary of Findings
    doc.add_page_break()
    doc.add_heading('SUMMARY OF FINDINGS', level=1)
    
    # Create comprehensive summary table
    summary_table = doc.add_table(rows=6, cols=2)
    summary_table.style = 'Table Grid'
    
    total_all_losses = past_total + future_total + data.life_care_plan_cost + data.household_services_cost
    
    summary_data = [
        ('Past Economic Losses', f"${past_total:,.2f}"),
        ('Future Economic Losses', f"${future_total:,.2f}"),
        ('Life Care Plan Costs', f"${data.life_care_plan_cost:,.2f}"),
        ('Household Services Loss', f"${data.household_services_cost:,.2f}"),
        ('', ''),
        ('TOTAL ECONOMIC LOSS', f"${total_all_losses:,.2f}")
    ]
    
    for i, (label, value) in enumerate(summary_data):
        summary_table.cell(i, 0).text = label
        summary_table.cell(i, 1).text = value
        if 'TOTAL' in label:
            # Make total row bold
            for paragraph in summary_table.cell(i, 0).paragraphs:
                for run in paragraph.runs:
                    run.bold = True
            for paragraph in summary_table.cell(i, 1).paragraphs:
                for run in paragraph.runs:
                    run.bold = True
    
    # Important Considerations
    doc.add_heading('Important Considerations', level=2)
    
    considerations_text = f"""
METHODOLOGY VALIDATION:
• All calculations employ standard economic principles and established methodologies
• Present value calculations utilize appropriate discount rates reflecting current economic conditions
• Demographic projections based on standard actuarial tables and work-life expectancy data
• Adjusted Earnings Factor methodology accounts for real-world economic conditions
• All projections subject to inherent uncertainty of future economic conditions

CALCULATION RELIABILITY:
• Economic assumptions based on historical data and reasonable projections
• All calculations verified for mathematical accuracy
• Results presented to a reasonable degree of economic certainty
• Sensitivity analysis available if required for key assumptions

PROFESSIONAL STANDARDS:
• Analysis follows accepted practices in forensic economics
• Methodology consistent with standards for expert economic testimony
• All assumptions and data sources documented and transparent
• Results comply with professional ethical standards for economic analysis

This analysis provides a comprehensive evaluation of the economic impact of {data.client_name}'s injuries, quantifying both the direct and indirect financial consequences of the incident. The methodology employed is scientifically sound, professionally accepted, and legally admissible.
"""
    doc.add_paragraph(considerations_text)
    
    # Conclusion section
    doc.add_page_break()
    doc.add_heading('CONCLUSION', level=1)
    
    conclusion_text = f"""
Based on this comprehensive economic analysis, {data.client_name} has sustained substantial economic losses as a direct result of the injuries sustained on {data.date_of_injury.strftime('%B %d, %Y')}.

TOTAL ECONOMIC IMPACT: ${total_all_losses:,.2f}

This figure represents the present value of all quantifiable economic losses, including:
• Lost earning capacity over the remaining work-life period
• Future medical and care costs related to the injury
• Loss of household service capacity
• All amounts calculated using established economic methodologies

The analysis employs reliable economic methods and data, providing a comprehensive assessment that quantifies the financial impact in present value terms. The methodologies used are consistent with those employed by forensic economists nationwide and have been tested for reasonableness and accuracy.

It is my professional opinion, based on a reasonable degree of economic certainty, that the totals presented fairly and accurately represent the economic damages sustained as a result of this incident. All calculations are documented, transparent, and consistent with accepted practices in forensic economics.

The economic impact extends beyond immediate medical costs to encompass the full spectrum of financial consequences, including diminished earning capacity, ongoing care requirements, and reduced capacity for household contributions. This comprehensive analysis provides the foundation for understanding the true economic magnitude of the losses sustained.
"""
    doc.add_paragraph(conclusion_text)
    
    # References section
    doc.add_page_break()
    doc.add_heading('REFERENCES', level=1)
    
    references_text = """
The following sources and references were utilized in this economic analysis:

• Tinari, F.D. (2016). "Demonstrating Lost Earnings: Algebraic vs. Spreadsheet Method." The Earnings Analyst, 15, 21–32. (Explains the algebraic approach to lost earnings calculations and its advantages in clarity)

• Tinari, F.D. (1998). "Household Services: Toward a More Comprehensive Measure." Journal of Forensic Economics, 11(3), 253–265. (Discusses inclusion of companionship and guidance in household services valuation)

• U.S. Bureau of Labor Statistics - Wage data, Consumer Price Index, and Occupational Employment Statistics for replacement cost methodology

• U.S. CDC Life Tables - Standard life expectancy data for demographic projections

• American Time Use Survey (ATUS) - Data on household service time allocation by demographics

• Zauner, H.A., Kozlowski, K., & Rella, J. (2025). "The Economist's Role in Forecasting Medical Costs in Personal Injury Disputes." Weaver Insights. (Methodology for projecting life care plan costs and medical inflation)

• The Knowles Group (2023). "Calculating Damages for Loss of Household Services." (Methods and data sources for household services valuation, including time-use data and replacement cost approach)

• Standard worklife expectancy tables adjusted for labor force participation rates

• National Association of Forensic Economics - Professional standards and methodologies

• Peer-reviewed economic literature on present value calculations, discount rates, and forensic economic methodology

All methodologies employed are consistent with established practices in forensic economics and have been subjected to peer review in academic literature. The principles and methods have been reliably applied to the specific facts of this case in accordance with Daubert standards for expert testimony.
"""
    doc.add_paragraph(references_text)
    
    # Glossary section
    doc.add_page_break()
    doc.add_heading('GLOSSARY OF TERMS', level=1)
    
    glossary_text = """
• Present Value (PV): The amount of money needed today to equal a future stream of payments, considering interest and inflation. For example, if you need $1,000 in five years and you can earn interest at 3%, the present value is about $863 today.

• Discount Rate: The interest rate used to convert future values into present value. It reflects the rate of return on safe investments (often based on government bonds for damage calculations). A higher discount rate means a lower present value for future amounts.

• Inflation Rate: The rate at which prices increase over time. Medical inflation refers to the rate of increase in healthcare costs, which historically has been higher than general inflation. We use inflation rates to project how wages or medical costs will grow in the future.

• Worklife Expectancy: The expected number of years a person will be active in the labor force, given their age, gender, and other factors. It accounts for probabilities of employment, unemployment, retirement, and mortality.

• Life Expectancy: The average number of additional years a person is expected to live, based on mortality tables. It provides the time frame for projecting future losses like medical costs or household services.

• Fringe Benefits: Non-wage compensation from employment, such as employer-paid health insurance, retirement contributions, bonuses, and other benefits. These are added to salary to reflect total compensation lost.

• Personal Consumption: In wrongful death economics, the portion of the decedent's income they would have spent on themselves. This is deducted because that portion is not a loss to surviving dependents.

• Life Care Plan: A document usually prepared by a medical or rehabilitation expert outlining an injured person's future medical and supportive care needs, along with the frequency and cost of each item or service.

• Replacement Cost Method: A way to value lost services by determining how much it would cost to replace those services in the market. For household services, it uses market wage rates for domestic tasks.

• Daubert Standard: The legal standard from Daubert v. Merrell Dow Pharmaceuticals (1993) used to evaluate whether an expert's testimony is admissible, focusing on the reliability and relevance of the methods used.

• Reasonable Degree of Economic Certainty: A phrase indicating that the expert's opinions are given with a high level of confidence, based on known data and sound methodology. It doesn't mean absolute certainty, but that the conclusions are well-founded.

• Adjusted Earnings Factor (AEF): A comprehensive methodology that applies multiple economic adjustments to gross earnings, including unemployment probability, tax rates, fringe benefits, and other factors to determine net economic loss.
"""
    doc.add_paragraph(glossary_text)
    
    # Professional statement
    doc.add_page_break()
    doc.add_heading('STATEMENT OF ETHICAL PRINCIPLES AND PRINCIPLES OF PROFESSIONAL PRACTICE', level=1)
    
    ethics_text = f"""
This economic analysis has been prepared in accordance with the ethical principles and professional practice standards of forensic economics. The analysis is based on reliable data, established methodologies, and reasonable assumptions. All calculations have been performed with due care and professional competence.

PROFESSIONAL STANDARDS:
• Independence and objectivity maintained in all professional work
• Adherence to highest standards of professional conduct
• Analysis based on reliable data and established methodologies
• Reasonable assumptions clearly documented and supported
• Calculations performed with due care and professional competence

RELIABILITY AND ACCURACY:
• All methodologies consistent with accepted practices in forensic economics
• Calculations verified for mathematical accuracy and reasonableness
• Data sources documented and assumptions clearly stated
• Results presented to a reasonable degree of economic certainty

ETHICAL COMPLIANCE:
• Complete and unbiased analysis of all relevant economic factors
• No material information omitted that would alter conclusions
• Independence from outcome maintained throughout analysis
• Professional duty to the court and truth upheld

The opinions expressed in this report are offered to a reasonable degree of economic certainty and are based on the information available at the time of preparation. This report contains the complete opinions of the analyst on the matters addressed herein. No part of the compensation for this analysis is contingent upon the results obtained or testimony given.

Respectfully submitted,

_________________________________
Forensic Economist
Date: {data.date_of_report.strftime('%B %d, %Y')}

Certification: I certify that I have applied the principles and methods reliably to the facts of this case, and I have not omitted any material information that would alter the conclusions. I am prepared to testify to these findings and explain the basis for each component of the loss analysis.
"""
    doc.add_paragraph(ethics_text)
    
    # Save to bytes
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io.getvalue()

# Streamlit Interface
def main():
    # Quick start section at the top
    st.info("🚀 **Quick Start:** Click '🔄 Load Sample Case' in the sidebar to test with comprehensive sample data, then go to the 'Generate Report' tab to create your professional economic loss report!")
    
    # Sidebar for main inputs
    with st.sidebar:
        st.header("📋 Report Configuration")
        
        # Sample case loader
        st.subheader("🚀 Quick Start")
        st.markdown("**Test the complete professional template:**")
        if st.button("🔄 Load Sample Case", type="primary", key="load_sample"):
            load_sample_case()
        st.caption("Loads: FIRST LAST - Mechanical Engineer with L4-L5 injury case")
        
        st.markdown("---")
        
        # Load/Save functionality
        st.subheader("Data Management")
        
        col1, col2 = st.columns(2)
        with col1:
            if st.button("💾 Save Data", key="save_data"):
                if st.session_state.report_data:
                    # Convert dates to strings for JSON
                    save_data = {}
                    for key, value in st.session_state.report_data.items():
                        if isinstance(value, date):
                            save_data[key] = value.isoformat()
                        else:
                            save_data[key] = value
                    
                    data_json = json.dumps(save_data, indent=2)
                    st.download_button(
                        "📥 Download JSON",
                        data_json,
                        f"economic_report_data_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                        "application/json",
                        key="download_json"
                    )
                else:
                    st.warning("No data to save. Please fill in the form first.")
        
        with col2:
            uploaded_file = st.file_uploader("📁 Load Data", type=['json'], key="upload_data")
            if uploaded_file:
                try:
                    loaded_data = json.load(uploaded_file)
                    # Convert date strings back to date objects
                    for key, value in loaded_data.items():
                        if key in ['date_of_birth', 'date_of_injury', 'date_of_report'] and isinstance(value, str):
                            try:
                                loaded_data[key] = datetime.fromisoformat(value).date()
                            except:
                                pass
                    st.session_state.report_data.update(loaded_data)
                    st.success("Data loaded successfully!")
                except Exception as e:
                    st.error(f"Error loading file: {str(e)}")
    
    # Main content area with tabs
    tab1, tab2, tab3, tab4, tab5 = st.tabs([
        "👤 Personal Info", 
        "💼 Employment", 
        "🏥 Injury Details", 
        "📊 Economic Factors", 
        "📋 Generate Report"
    ])
    
    # Initialize data object
    data = EconomicData()
    
    with tab1:
        st.header("Personal Information")
        
        col1, col2 = st.columns(2)
        with col1:
            data.client_name = st.text_input(
                "Client Name", 
                value=st.session_state.report_data.get('client_name', ''),
                key='client_name',
                help=FIELD_HELP['client_name']
            )
            data.date_of_birth = st.date_input(
                "Date of Birth", 
                value=st.session_state.report_data.get('date_of_birth', date(1980, 1, 1)),
                key='date_of_birth',
                help=FIELD_HELP['date_of_birth']
            )
            data.gender = st.selectbox(
                "Gender", 
                ["Male", "Female"],
                index=0 if st.session_state.report_data.get('gender', 'Male') == 'Male' else 1,
                key='gender',
                help=FIELD_HELP['gender']
            )
            data.education = st.text_input(
                "Education Level", 
                value=st.session_state.report_data.get('education', ''),
                key='education',
                help=FIELD_HELP['education']
            )
        
        with col2:
            data.date_of_injury = st.date_input(
                "Date of Injury", 
                value=st.session_state.report_data.get('date_of_injury', date(2023, 1, 1)),
                key='date_of_injury',
                help=FIELD_HELP['date_of_injury']
            )
            data.date_of_report = st.date_input(
                "Date of Report", 
                value=st.session_state.report_data.get('date_of_report', date.today()),
                key='date_of_report',
                help=FIELD_HELP['date_of_report']
            )
            data.marital_status = st.text_input(
                "Marital Status", 
                value=st.session_state.report_data.get('marital_status', ''),
                key='marital_status',
                help=FIELD_HELP['marital_status']
            )
        
        # Calculate age at injury
        data.age_at_injury = calculate_age(data.date_of_birth, data.date_of_injury)
        st.info(f"Age at injury: {data.age_at_injury:.1f} years")
        
        st.subheader("Geographic Information")
        col3, col4, col5 = st.columns(3)
        with col3:
            data.residence_state = st.text_input(
                "State of Residence", 
                value=st.session_state.report_data.get('residence_state', ''),
                key='residence_state',
                help=FIELD_HELP['residence_state']
            )
        with col4:
            data.residence_county = st.text_input(
                "County", 
                value=st.session_state.report_data.get('residence_county', ''),
                key='residence_county',
                help=FIELD_HELP['residence_county']
            )
        with col5:
            data.msa = st.text_input(
                "MSA (Metropolitan Statistical Area)", 
                value=st.session_state.report_data.get('msa', ''),
                key='msa',
                help=FIELD_HELP['msa']
            )
    
    with tab2:
        st.header("Employment Information")
        
        col1, col2 = st.columns(2)
        with col1:
            data.pre_injury_occupation = st.text_input(
                "Pre-Injury Occupation", 
                value=st.session_state.report_data.get('pre_injury_occupation', ''),
                key='pre_injury_occupation',
                help=FIELD_HELP['pre_injury_occupation']
            )
            data.pre_injury_annual_income = st.number_input(
                "Pre-Injury Annual Income ($)", 
                min_value=0.0, 
                value=float(st.session_state.report_data.get('pre_injury_annual_income', 50000.0)),
                step=1000.0,
                key='pre_injury_annual_income',
                help=FIELD_HELP['pre_injury_annual_income']
            )
            data.employer = st.text_input(
                "Employer", 
                value=st.session_state.report_data.get('employer', ''),
                key='employer',
                help=FIELD_HELP['employer']
            )
        
        with col2:
            data.pre_injury_hourly_rate = st.number_input(
                "Pre-Injury Hourly Rate ($)", 
                min_value=0.0, 
                value=float(st.session_state.report_data.get('pre_injury_hourly_rate', 24.04)),
                step=0.50,
                key='pre_injury_hourly_rate',
                help=FIELD_HELP['pre_injury_hourly_rate']
            )
            data.years_with_employer = st.number_input(
                "Years with Employer", 
                min_value=0.0, 
                value=float(st.session_state.report_data.get('years_with_employer', 5.0)),
                step=0.5,
                key='years_with_employer',
                help=FIELD_HELP['years_with_employer']
            )
        
        st.subheader("Post-Injury Employment")
        col3, col4 = st.columns(2)
        with col3:
            data.post_injury_occupation = st.text_input(
                "Post-Injury Occupation", 
                value=st.session_state.report_data.get('post_injury_occupation', ''),
                key='post_injury_occupation',
                help=FIELD_HELP['post_injury_occupation']
            )
            data.post_injury_capacity = st.text_area(
                "Post-Injury Work Capacity", 
                value=st.session_state.report_data.get('post_injury_capacity', ''),
                key='post_injury_capacity',
                help=FIELD_HELP['post_injury_capacity']
            )
        
        with col4:
            data.post_injury_annual_income = st.number_input(
                "Post-Injury Annual Income ($)", 
                min_value=0.0, 
                value=float(st.session_state.report_data.get('post_injury_annual_income', 0.0)),
                step=1000.0,
                key='post_injury_annual_income',
                help=FIELD_HELP['post_injury_annual_income']
            )
            data.labor_market_reduction = st.number_input(
                "Labor Market Reduction (%)", 
                min_value=0.0,
                max_value=100.0,
                value=float(st.session_state.report_data.get('labor_market_reduction', 0.0)),
                step=1.0,
                key='labor_market_reduction',
                help=FIELD_HELP['labor_market_reduction']
            ) / 100
    
    with tab3:
        st.header("Injury and Medical Information")
        
        col1, col2 = st.columns(2)
        with col1:
            data.injury_description = st.text_area(
                "Injury Description", 
                value=st.session_state.report_data.get('injury_description', ''),
                key='injury_description',
                height=100,
                help=FIELD_HELP['injury_description']
            )
            data.body_parts_injured = st.text_input(
                "Body Parts Injured", 
                value=st.session_state.report_data.get('body_parts_injured', ''),
                key='body_parts_injured',
                help=FIELD_HELP['body_parts_injured']
            )
        
        with col2:
            data.medical_treatment = st.text_area(
                "Medical Treatment", 
                value=st.session_state.report_data.get('medical_treatment', ''),
                key='medical_treatment',
                height=100,
                help=FIELD_HELP['medical_treatment']
            )
            data.current_medical_status = st.text_area(
                "Current Medical Status", 
                value=st.session_state.report_data.get('current_medical_status', ''),
                key='current_medical_status',
                height=100,
                help=FIELD_HELP['current_medical_status']
            )
    
    with tab4:
        st.header("Economic Factors and Life Expectancy")
        
        st.subheader("Life Expectancy Data")
        col1, col2, col3 = st.columns(3)
        with col1:
            data.life_expectancy = st.number_input(
                "Life Expectancy (years)", 
                min_value=0.0,
                max_value=120.0,
                value=float(st.session_state.report_data.get('life_expectancy', 78.5)),
                step=0.1,
                key='life_expectancy',
                help=FIELD_HELP['life_expectancy']
            )
        with col2:
            data.work_life_expectancy = st.number_input(
                "Work Life Expectancy (years)", 
                min_value=0.0,
                max_value=80.0,
                value=float(st.session_state.report_data.get('work_life_expectancy', 45.0)),
                step=0.1,
                key='work_life_expectancy',
                help=FIELD_HELP['work_life_expectancy']
            )
        with col3:
            data.years_to_final_separation = st.number_input(
                "Years to Final Separation", 
                min_value=0.0,
                value=float(st.session_state.report_data.get('years_to_final_separation', 0.0)),
                step=0.1,
                key='years_to_final_separation',
                help=FIELD_HELP['years_to_final_separation']
            )
        
        # Calculate and display statistical dates
        retirement_date, death_date = calculate_statistical_dates(
            data.date_of_injury, data.life_expectancy, data.work_life_expectancy
        )
        
        col4, col5 = st.columns(2)
        with col4:
            st.info(f"Statistical Retirement Date: {retirement_date.strftime('%m/%d/%Y')}")
        with col5:
            st.info(f"Statistical Death Date: {death_date.strftime('%m/%d/%Y')}")
        
        st.subheader("Economic Growth and Discount Rates")
        col6, col7 = st.columns(2)
        with col6:
            data.wage_growth_rate = st.number_input(
                "Wage Growth Rate (%)", 
                min_value=0.0,
                max_value=20.0,
                value=float(st.session_state.report_data.get('wage_growth_rate', 3.0)),
                step=0.1,
                key='wage_growth_rate',
                help=FIELD_HELP['wage_growth_rate']
            ) / 100
        with col7:
            data.discount_rate = st.number_input(
                "Discount Rate (%)", 
                min_value=0.0,
                max_value=20.0,
                value=float(st.session_state.report_data.get('discount_rate', 4.0)),
                step=0.1,
                key='discount_rate',
                help=FIELD_HELP['discount_rate']
            ) / 100
        
        st.subheader("Adjustment Factors")
        col8, col9 = st.columns(2)
        with col8:
            data.unemployment_rate = st.number_input(
                "Unemployment Rate (%)", 
                min_value=0.0,
                max_value=30.0,
                value=float(st.session_state.report_data.get('unemployment_rate', 3.5)),
                step=0.1,
                key='unemployment_rate',
                help=FIELD_HELP['unemployment_rate']
            ) / 100
            data.tax_rate = st.number_input(
                "Tax Rate (%)", 
                min_value=0.0,
                max_value=50.0,
                value=float(st.session_state.report_data.get('tax_rate', 12.0)),
                step=0.1,
                key='tax_rate',
                help=FIELD_HELP['tax_rate']
            ) / 100
        with col9:
            data.fringe_benefits_rate = st.number_input(
                "Fringe Benefits Rate (%)", 
                min_value=0.0,
                max_value=50.0,
                value=float(st.session_state.report_data.get('fringe_benefits_rate', 6.0)),
                step=0.1,
                key='fringe_benefits_rate',
                help=FIELD_HELP['fringe_benefits_rate']
            ) / 100
            data.personal_consumption_rate = st.number_input(
                "Personal Consumption Rate (%)", 
                min_value=0.0,
                max_value=50.0,
                value=float(st.session_state.report_data.get('personal_consumption_rate', 0.0)),
                step=0.1,
                key='personal_consumption_rate',
                help=FIELD_HELP['personal_consumption_rate'] + " - Use for wrongful death cases"
            ) / 100
        
        st.subheader("Additional Costs")
        col10, col11 = st.columns(2)
        with col10:
            data.life_care_plan_cost = st.number_input(
                "Life Care Plan Cost ($)", 
                min_value=0.0,
                value=float(st.session_state.report_data.get('life_care_plan_cost', 0.0)),
                step=1000.0,
                key='life_care_plan_cost',
                help=FIELD_HELP['life_care_plan_cost']
            )
        with col11:
            data.household_services_cost = st.number_input(
                "Household Services Cost ($)", 
                min_value=0.0,
                value=float(st.session_state.report_data.get('household_services_cost', 0.0)),
                step=1000.0,
                key='household_services_cost',
                help=FIELD_HELP['household_services_cost']
            )
    
    with tab5:
        st.header("Generate Economic Loss Report")
        
        # Testing instructions
        if not any(st.session_state.report_data.values()):
            st.warning("⚠️ **No case data loaded.** Click '🔄 Load Sample Case' in the sidebar to test with comprehensive sample data!")
        else:
            st.success("✅ **Case data loaded:** " + st.session_state.report_data.get('client_name', 'Unknown Client'))
        
        # Update session state with current data
        for key, value in data.__dict__.items():
            st.session_state.report_data[key] = value
        
        # Calculate button
        if st.button("🔄 Calculate Economic Losses", type="primary", key="calculate_button"):
            with st.spinner("Calculating economic losses..."):
                try:
                    # Generate calculations
                    past_df, future_df, aef_result, retirement_date, death_date = generate_economic_schedule(data)
                    
                    # Store results in session state
                    st.session_state.calculation_results = {
                        'past_df': past_df,
                        'future_df': future_df,
                        'aef_result': aef_result,
                        'retirement_date': retirement_date,
                        'death_date': death_date,
                        'data': data
                    }
                    st.session_state.show_results = True
                    
                    st.success("✅ Calculations completed successfully!")
                    
                except Exception as e:
                    st.error(f"Error in calculations: {str(e)}")
                    st.exception(e)
        
        # Display results if available
        if st.session_state.show_results and st.session_state.calculation_results:
            results = st.session_state.calculation_results
            past_df = results['past_df']
            future_df = results['future_df']
            aef_result = results['aef_result']
            retirement_date = results['retirement_date']
            death_date = results['death_date']
            data = results['data']
            
            # Display results
            st.subheader("📊 Calculation Results")
            
            # Summary metrics
            past_total = past_df['Present Value'].sum() if not past_df.empty else 0
            future_total = future_df['Present Value'].sum() if not future_df.empty else 0
            total_loss = past_total + future_total + data.life_care_plan_cost
            
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.metric("Past Losses", f"${past_total:,.2f}")
            with col2:
                st.metric("Future Losses", f"${future_total:,.2f}")
            with col3:
                st.metric("Life Care Plan", f"${data.life_care_plan_cost:,.2f}")
            with col4:
                st.metric("Total Economic Loss", f"${total_loss:,.2f}")
            
            # AEF breakdown
            st.subheader("🧮 Adjusted Earnings Factor (AEF)")
            aef_data = []
            for factor_name, factor_value, cumulative in aef_result['factors']:
                aef_data.append({
                    'Factor': factor_name,
                    'Value': f"{factor_value:.3f}",
                    'Percentage': f"{factor_value * 100:.1f}%",
                    'Cumulative': f"{cumulative:.3f}"
                })
            
            aef_df = pd.DataFrame(aef_data)
            st.dataframe(aef_df, use_container_width=True, hide_index=True)
            
            # Past losses table
            if not past_df.empty:
                st.subheader("📅 Past Economic Losses")
                st.dataframe(past_df, use_container_width=True, hide_index=True)
            
            # Future losses table
            if not future_df.empty:
                st.subheader("🔮 Future Economic Losses")
                st.dataframe(future_df, use_container_width=True, hide_index=True)
            
            # Generate Word report section
            st.subheader("📄 Download Report")
            
            col1, col2 = st.columns([1, 1])
            
            with col1:
                if st.button("📄 Generate Complete Professional Report", type="primary", key="generate_word"):
                    with st.spinner("Generating comprehensive professional Word document..."):
                        try:
                            doc_bytes = create_complete_professional_report(data, past_df, future_df, aef_result, retirement_date, death_date)
                            
                            if doc_bytes:
                                # Create download button
                                file_name = f"Economic_Loss_Report_{data.client_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"
                                
                                st.download_button(
                                    label="📥 Download Complete Professional Report",
                                    data=doc_bytes,
                                    file_name=file_name,
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key="word_download_final"
                                )
                                
                                st.success(f"✅ Word report generated successfully! ({len(doc_bytes):,} bytes)")
                            else:
                                st.error("Failed to generate Word document")
                        except Exception as e:
                            st.error(f"Error generating Word report: {str(e)}")
                            st.exception(e)
            
            with col2:
                st.info("💡 The Complete Professional Report includes:\n- Kincaid Wolstein professional branding\n- Daubert-compliant methodology\n- Comprehensive expert analysis\n- Past & future loss calculations\n- Medical cost projections\n- Household services analysis\n- Expert qualifications\n- References and citations\n- Lay-person explanations\n- All template sections from pages 3-16")
            
            # Export data options
            st.subheader("📊 Export Data")
            col1, col2, col3 = st.columns(3)
            
            with col1:
                if not past_df.empty:
                    csv_past = past_df.to_csv(index=False)
                    st.download_button(
                        "📊 Past Losses CSV",
                        csv_past,
                        f"past_losses_{data.client_name.replace(' ', '_')}.csv",
                        "text/csv",
                        key="csv_past_download"
                    )
            
            with col2:
                if not future_df.empty:
                    csv_future = future_df.to_csv(index=False)
                    st.download_button(
                        "📊 Future Losses CSV",
                        csv_future,
                        f"future_losses_{data.client_name.replace(' ', '_')}.csv",
                        "text/csv",
                        key="csv_future_download"
                    )
            
            with col3:
                # Combined Excel export
                excel_io = io.BytesIO()
                try:
                    with pd.ExcelWriter(excel_io, engine='openpyxl') as writer:
                        if not past_df.empty:
                            past_df.to_excel(writer, sheet_name='Past Losses', index=False)
                        if not future_df.empty:
                            future_df.to_excel(writer, sheet_name='Future Losses', index=False)
                        aef_df.to_excel(writer, sheet_name='AEF Breakdown', index=False)
                    
                    excel_data = excel_io.getvalue()
                    st.download_button(
                        "📊 Complete Excel Report",
                        excel_data,
                        f"economic_analysis_{data.client_name.replace(' ', '_')}.xlsx",
                        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        key="excel_download"
                    )
                except Exception as e:
                    st.error(f"Error creating Excel file: {str(e)}")

if __name__ == "__main__":
    main()