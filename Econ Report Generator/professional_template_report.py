"""
Professional Economic Loss Report Generator
Based on Kincaid Wolstein Template Structure
"""

from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from datetime import date, datetime
import pandas as pd
from dataclasses import dataclass

def create_professional_report(data, past_df: pd.DataFrame, future_df: pd.DataFrame, 
                             aef_result: dict, retirement_date: date, death_date: date) -> bytes:
    """Generate professional economic loss report matching exact template format"""
    
    doc = Document()
    
    # Calculate totals
    past_total = past_df['Present Value'].sum() if not past_df.empty and 'Present Value' in past_df.columns else 0
    future_total = future_df['Present Value'].sum() if not future_df.empty and 'Present Value' in future_df.columns else 0
    total_economic_loss = past_total + future_total + data.life_care_plan_cost + data.household_services_cost
    
    # Set document margins
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    
    # TITLE PAGE
    # Company logo and header section
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Add the Kincaid Wolstein logo styling
    title_run = header_para.add_run("Kincaid\nWolstein")
    title_run.font.size = Inches(0.4)
    title_run.font.color.rgb = RGBColor(52, 143, 226)  # Blue color
    title_run.bold = True
    
    subtitle_run = header_para.add_run("\nVOCATIONAL &\nREHABILITATION\nSERVICES")
    subtitle_run.font.size = Inches(0.15)
    
    # Draft privilege notice
    privilege_para = doc.add_paragraph()
    privilege_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    privilege_run = privilege_para.add_run("DRAFT-WORK PRODUCT PRIVILEGE")
    privilege_run.font.color.rgb = RGBColor(52, 143, 226)
    
    # Add spacing
    for _ in range(6):
        doc.add_paragraph()
    
    # Main title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("APPRAISAL OF ECONOMIC LOSS\nRESULTING FROM INJURY TO ")
    title_run.font.size = Inches(0.2)
    title_run.bold = True
    
    name_run = title_para.add_run(data.client_name.upper())
    name_run.font.size = Inches(0.2)
    name_run.bold = True
    name_run.font.color.rgb = RGBColor(52, 143, 226)
    
    # Add spacing
    for _ in range(4):
        doc.add_paragraph()
    
    # Prepared by section
    prepared_para = doc.add_paragraph()
    prepared_para.add_run("PREPARED BY:\t").bold = True
    prepared_para.add_run("Kincaid Wolstein Vocational and Rehabilitation Services\n")
    prepared_para.add_run("\t\t\tOne University Plaza ~ Suite 302\n")
    prepared_para.add_run("\t\t\tHackensack, New Jersey 07601\n")
    prepared_para.add_run("\t\t\tPhone: (201) 343-0700\n")
    prepared_para.add_run("\t\t\tFax: (201) 343-0757\n")
    
    doc.add_paragraph()
    
    # Case details
    prepared_for_para = doc.add_paragraph()
    prepared_for_para.add_run("PREPARED FOR:\n").bold = True
    
    doc.add_paragraph()
    
    regarding_para = doc.add_paragraph()
    regarding_para.add_run("REGARDING:\n").bold = True
    
    doc.add_paragraph()
    
    dob_para = doc.add_paragraph()
    dob_para.add_run("DATE OF BIRTH:\n").bold = True
    
    doc.add_paragraph()
    
    report_date_para = doc.add_paragraph()
    report_date_para.add_run("REPORT DATE:\n").bold = True
    
    doc.add_page_break()
    
    # TABLE OF CONTENTS
    toc_title = doc.add_heading('Table of Contents', level=1)
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    toc_items = [
        "CERTIFICATION.............................................................................................................3",
        "PURPOSE OF APPRAISAL .............................................................................................4", 
        "OPINION OF ECONOMIC LOSSES................................................................................4",
        "BACKGROUND FACTS AND ASSUMPTIONS ...............................................................5",
        "    Summary Information..................................................................................................5",
        "    Life Expectancy...........................................................................................................5",
        "    Statistical retirement age ............................................................................................5",
        "    Expected working years .............................................................................................5",
        "    Worklife-to Retirement Ratio ......................................................................................6",
        "    Description of Incident ................................................................................................6",
        "    Occupation and Employment......................................................................................6",
        "    Earnings History..........................................................................................................7",
        "    Pre-injury Earnings Capacity: .....................................................................................7",
        "    Fringe Benefits............................................................................................................7",
        "    Functionality and Future Employability ......................................................................8",
        "    Post-injury Earnings Capacity ....................................................................................9",
        "    Growth Wage Rates....................................................................................................9",
        "    Household Services ..................................................................................................10",
        "COMPONENTS OF ANALYSIS .....................................................................................12",
        "PRE‐INJURY ADJUSTED EARNINGS .........................................................................13",
        "PRE‐INJURY ADJUSTED EARNINGS FOLLOWING RTW .........................................15",
        "POST‐INJURY ADJUSTED EARNINGS FOLLOWING DEPARTURE FROM FULL",
        "DUTY POSITION..........................................................................................................16",
        "COST OF LIFETIME CARE ..........................................................................................18",
        "    Methodology .............................................................................................................20",
        "    Summary of Findings ................................................................................................20",
        "    Important Considerations..........................................................................................21",
        "COST OF LIFETIME CARE ..........................................................................................23",
        "TABLES ........................................................................................................................24",
        "STATEMENT OF ETHICAL PRINCIPLES AND PRINCIPLES OF PROFESSIONAL ...38",
        "PRACTICE ....................................................................................................................38"
    ]
    
    for item in toc_items:
        doc.add_paragraph(item)
    
    doc.add_page_break()
    
    # ECONOMIC LOSS APPRAISAL REPORT
    doc.add_heading('ECONOMIC LOSS APPRAISAL REPORT', level=1)
    
    # Introduction
    doc.add_heading('Introduction', level=2)
    intro_text = f"""This Economic Loss Appraisal Report has been prepared by [Expert Name], a forensic economist, in the matter of [Case Name], to evaluate the economic losses resulting from {data.client_name}'s injuries (or wrongful death). The purpose of this report is to quantify various categories of economic damages in a manner that is clear, comprehensive, and compliant with the Daubert standard for expert testimony. The key components of loss addressed include:

• Lost Wages/Earnings Capacity – the value of past and future income (and benefits) lost due to the incident.
• Future Medical and Healthcare Costs – the present value of reasonable future medical expenses related to the injury.
• Loss of Household Services – the economic value of household tasks and services {data.client_name} can no longer perform.

All findings are presented in plain language with technical details explained for a general audience. Specialized terms are defined throughout the report and in a glossary. The methodologies used are grounded in well-established economic principles and reliable methods that have been published in peer-reviewed literature and are generally accepted in the field. This report adheres to Daubert criteria by using reliable principles and methods, referencing peer-reviewed sources, discussing potential error rates or uncertainties, and applying the methods to the facts of this case. All assumptions, data sources, and calculations are documented in the sections below, and all opinions are stated to a reasonable degree of economic certainty.

Report Structure: After summarizing the data and assumptions considered, the report details each category of loss in turn (Lost Wages, Future Medical Costs, Household Services), followed by a summary of total losses. Each section outlines the methodology and reasoning, ensuring transparency and layperson accessibility. A final section provides requisite expert disclosures, including qualifications, materials reviewed, assumptions, exhibits, and a signature attestation."""
    
    doc.add_paragraph(intro_text)
    
    # Materials Considered
    doc.add_heading('Materials Considered', level=2)
    materials_text = f"""In preparing this analysis, I have reviewed and relied upon the following materials and data (among others):

• Case Documents: [List key documents: e.g. Complaint, deposition of {data.client_name}, accident reports].
• Medical Records & Life Care Plan: [e.g. Medical reports from Dr. A; Life Care Plan by [Name] dated ___ outlining future care needs].
• Employment & Earnings Records: [e.g. pay stubs, W-2 forms, tax returns, employment file from {data.client_name}'s employer].
• Economic & Statistical Data: U.S. Bureau of Labor Statistics (wage data, Consumer Price Index), published worklife expectancy tables, life expectancy tables (e.g. U.S. CDC Life Tables), and relevant economic research literature.
• Other: [Any other data: e.g. vocational expert report, family testimony on household services, etc.].

(Modify the above list as needed for the specific case, ensuring all materials considered are listed.)"""
    
    doc.add_paragraph(materials_text)
    
    # Key Assumptions
    doc.add_heading('Key Assumptions', level=2)
    assumptions_text = f"""The following assumptions underlie the calculations in this report (to be adjusted per case specifics):

• Employment but for Incident: It is assumed that absent the incident, {data.client_name} would have continued working in their usual capacity up to a normal retirement age of [X] years (or for [Y] additional years of worklife). This worklife expectancy is based on statistical averages adjusted for {data.client_name}'s age, gender, and work history.

• Post-Incident Work Capacity: {data.client_name} is now unable to work (or can only work in a limited capacity earning ${data.post_injury_annual_income:,.2f} per year), per medical/vocational evidence, resulting in a loss of earning capacity as detailed below.

• Fringe Benefits: Employer-provided benefits (health insurance, retirement contributions, etc.) comprised approximately {data.fringe_benefits_rate*100:.1f}% of wages and are included as part of the lost compensation.

• Income Taxes: Calculations of lost earnings are presented on an [after-tax basis] (if applicable under jurisdiction) so that the award reflects net take-home pay loss. A combined federal/state tax rate of {data.tax_rate*100:.1f}% is assumed for this purpose.

• Life Expectancy: {data.client_name} has a remaining life expectancy of {data.life_expectancy:.1f} years, based on [source, e.g. U.S. Life Tables], which is used to project future losses through the year {death_date.year}. If the injury is expected to impact longevity, medical expert input is used to adjust this expectation.

• Discount Rate: A discount rate of {data.discount_rate*100:.1f}% is used to convert future dollars to present value. This rate is chosen to reflect a risk-free rate of return (e.g. based on U.S. Treasury or tax-free municipal bond yields) appropriate for a lump-sum award.

• Inflation Rates: Future wage growth and medical cost inflation are assumed at {data.wage_growth_rate*100:.1f}% annually (based on historical data and current economic forecasts for wages and healthcare costs, respectively). These inflation assumptions are paired with the discount rate to ensure consistency (e.g. real vs. nominal projections).

• Household Services: It is assumed that {data.client_name} performed approximately ___ hours per week of household and family services pre-incident (based on time-use data or family testimony), and that due to the injury [he/she] can no longer perform [all or a specific portion] of these tasks. The types of services affected include [list: e.g. cleaning, cooking, childcare, yard work].

• Mitigation: Any mitigation or offset (such as actual earnings post-incident, or replacement services provided by others) has been considered and will be noted in the calculations. It is assumed that {data.client_name} has made reasonable efforts to mitigate losses where possible.

(Each assumption should be reviewed and modified to fit the facts of the case. Additional assumptions can be added as needed.)"""
    
    doc.add_paragraph(assumptions_text)
    
    doc.add_page_break()
    
    # Continue with remaining sections following the exact template structure...
    # This would continue with all the detailed sections from the template
    
    # Save document
    import io
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io.getvalue()