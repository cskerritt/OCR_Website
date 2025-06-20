import streamlit as st
import pandas as pd
import numpy as np
from datetime import date, datetime, timedelta
from dataclasses import dataclass
import io
from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.shared import OxmlElement, qn
import json
import matplotlib.pyplot as plt

# Page configuration
st.set_page_config(
    page_title="Economic Report Generator", 
    layout="wide", 
    initial_sidebar_state="expanded"
)

st.title("📊 Economic Loss Report Generator")
st.markdown("Generate comprehensive economic loss reports with professional calculations and formatting.")

# Initialize session state for data persistence
if 'report_data' not in st.session_state:
    st.session_state.report_data = {}

@dataclass
class EconomicData:
    """Data structure for economic analysis"""
    # Personal Information
    client_name: str = ""
    date_of_birth: date = date(1980, 1, 1)
    date_of_injury: date = date(2023, 1, 1)
    date_of_report: date = date.today()
    age_at_injury: float = 0.0
    gender: str = "Male"
    education: str = ""
    marital_status: str = ""
    
    # Geographic Information
    residence_state: str = ""
    residence_county: str = ""
    msa: str = ""
    
    # Employment Information
    pre_injury_occupation: str = ""
    pre_injury_annual_income: float = 0.0
    pre_injury_hourly_rate: float = 0.0
    employer: str = ""
    years_with_employer: float = 0.0
    
    # Injury Information
    injury_description: str = ""
    body_parts_injured: str = ""
    medical_treatment: str = ""
    current_medical_status: str = ""
    
    # Work Life Expectancy Data
    life_expectancy: float = 78.5
    work_life_expectancy: float = 45.0
    years_to_final_separation: float = 0.0
    statistical_retirement_date: date = date.today()
    
    # Economic Factors
    wage_growth_rate: float = 0.03
    discount_rate: float = 0.04
    unemployment_rate: float = 0.035
    tax_rate: float = 0.12
    fringe_benefits_rate: float = 0.06
    personal_consumption_rate: float = 0.25
    
    # Post-Injury Information
    post_injury_capacity: str = ""
    post_injury_occupation: str = ""
    post_injury_annual_income: float = 0.0
    labor_market_reduction: float = 0.0
    
    # Additional Costs
    life_care_plan_cost: float = 0.0
    household_services_cost: float = 0.0

def calculate_age(birth_date: date, reference_date: date) -> float:
    """Calculate age in years with decimal precision"""
    return (reference_date - birth_date).days / 365.25

def calculate_statistical_dates(doi: date, le: float, wle: float) -> tuple:
    """Calculate statistical retirement and death dates"""
    retirement_date = doi + timedelta(days=int(wle * 365.25))
    death_date = doi + timedelta(days=int(le * 365.25))
    return retirement_date, death_date

def calculate_worklife_factor(wle: float, yfs: float) -> float:
    """Calculate worklife participation factor"""
    if yfs > 0:
        return wle / (wle + yfs)
    return 1.0

def calculate_adjusted_earnings_factor(unemployment: float, tax: float, 
                                     fringe: float = 0.0, consumption: float = 0.0,
                                     worklife_factor: float = 1.0) -> dict:
    """Calculate comprehensive adjusted earnings factor"""
    factors = []
    current_factor = 1.0
    
    # Worklife factor
    current_factor *= worklife_factor
    factors.append(("Worklife Factor", worklife_factor, current_factor))
    
    # Fringe benefits (additive)
    if fringe > 0:
        current_factor *= (1 + fringe)
        factors.append(("Fringe Benefits", 1 + fringe, current_factor))
    
    # Unemployment reduction
    current_factor *= (1 - unemployment)
    factors.append(("Unemployment Factor", 1 - unemployment, current_factor))
    
    # Tax reduction
    current_factor *= (1 - tax)
    factors.append(("Tax Factor", 1 - tax, current_factor))
    
    # Personal consumption (for wrongful death)
    if consumption > 0:
        current_factor = current_factor - consumption
        factors.append(("Personal Consumption", consumption, current_factor))
    
    return {
        'final_aef': current_factor,
        'factors': factors
    }

def generate_economic_schedule(data: EconomicData) -> tuple:
    """Generate past and future economic loss schedules"""
    
    # Calculate key dates and factors
    retirement_date, death_date = calculate_statistical_dates(
        data.date_of_injury, data.life_expectancy, data.work_life_expectancy
    )
    
    worklife_factor = calculate_worklife_factor(
        data.work_life_expectancy, data.years_to_final_separation
    )
    
    aef_result = calculate_adjusted_earnings_factor(
        data.unemployment_rate, data.tax_rate, 
        data.fringe_benefits_rate, data.personal_consumption_rate,
        worklife_factor
    )
    
    # Past losses (DOI to DOR)
    past_years = []
    report_year = data.date_of_report.year
    injury_year = data.date_of_injury.year
    
    for year in range(injury_year, report_year + 1):
        age = calculate_age(data.date_of_birth, date(year, 7, 1))
        years_since_injury = year - injury_year
        
        # Calculate portion of year
        if year == injury_year:
            # Partial year from injury date
            days_remaining = (date(year, 12, 31) - data.date_of_injury).days + 1
            portion = days_remaining / 365
        elif year == report_year:
            # Partial year to report date
            days_elapsed = data.date_of_report.timetuple().tm_yday
            portion = days_elapsed / 365
        else:
            portion = 1.0
        
        # Calculate earnings
        pre_injury_earnings = data.pre_injury_annual_income * ((1 + data.wage_growth_rate) ** years_since_injury)
        nominal_loss = pre_injury_earnings * portion
        adjusted_loss = nominal_loss * aef_result['final_aef']
        
        # Present value (discount back to injury date)
        pv_loss = adjusted_loss / ((1 + data.discount_rate) ** years_since_injury)
        
        past_years.append({
            'Year': year,
            'Age': round(age, 1),
            'Portion': round(portion * 100, 1),
            'Pre-Injury Earnings': round(pre_injury_earnings, 2),
            'Nominal Loss': round(nominal_loss, 2),
            'AEF': round(aef_result['final_aef'] * 100, 2),
            'Adjusted Loss': round(adjusted_loss, 2),
            'Present Value': round(pv_loss, 2)
        })
    
    # Future losses (DOR to retirement)
    future_years = []
    retirement_year = retirement_date.year
    
    for year in range(report_year, retirement_year + 1):
        if year == report_year:
            continue  # Already handled in past
            
        age = calculate_age(data.date_of_birth, date(year, 7, 1))
        years_since_injury = year - injury_year
        
        # Calculate portion of year
        if year == retirement_year:
            # Partial year to retirement
            retirement_day = retirement_date.timetuple().tm_yday
            portion = retirement_day / 365
        else:
            portion = 1.0
        
        # Calculate earnings
        pre_injury_earnings = data.pre_injury_annual_income * ((1 + data.wage_growth_rate) ** years_since_injury)
        post_injury_earnings = data.post_injury_annual_income * ((1 + data.wage_growth_rate) ** (years_since_injury - (report_year - injury_year)))
        
        nominal_loss = (pre_injury_earnings - post_injury_earnings) * portion
        adjusted_loss = nominal_loss * aef_result['final_aef']
        
        # Present value (discount to injury date)
        pv_loss = adjusted_loss / ((1 + data.discount_rate) ** years_since_injury)
        
        future_years.append({
            'Year': year,
            'Age': round(age, 1),
            'Portion': round(portion * 100, 1),
            'Pre-Injury Earnings': round(pre_injury_earnings, 2),
            'Post-Injury Earnings': round(post_injury_earnings, 2),
            'Nominal Loss': round(nominal_loss, 2),
            'AEF': round(aef_result['final_aef'] * 100, 2),
            'Adjusted Loss': round(adjusted_loss, 2),
            'Present Value': round(pv_loss, 2)
        })
    
    past_df = pd.DataFrame(past_years)
    future_df = pd.DataFrame(future_years)
    
    return past_df, future_df, aef_result, retirement_date, death_date

def create_word_report(data: EconomicData, past_df: pd.DataFrame, 
                      future_df: pd.DataFrame, aef_result: dict, 
                      retirement_date: date, death_date: date) -> bytes:
    """Generate comprehensive Word report"""
    
    doc = Document()
    
    # Set document margins and orientation
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    
    # Title page
    title = doc.add_heading('ECONOMIC LOSS ANALYSIS', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Client information
    client_info = doc.add_paragraph()
    client_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    client_info.add_run(f"RE: {data.client_name}").bold = True
    
    doc.add_paragraph()
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.add_run(f"Report Date: {data.date_of_report.strftime('%B %d, %Y')}")
    
    # Page break
    doc.add_page_break()
    
    # Executive Summary
    doc.add_heading('EXECUTIVE SUMMARY', level=1)
    
    summary_text = f"""
This report presents an economic analysis of the financial losses sustained by {data.client_name} 
as a result of injuries sustained on {data.date_of_injury.strftime('%B %d, %Y')}. The analysis 
considers both past and future economic losses through the individual's projected work life.

Key findings include:
• Total Past Losses: ${past_df['Present Value'].sum():,.2f}
• Total Future Losses: ${future_df['Present Value'].sum():,.2f}
• Total Economic Loss: ${(past_df['Present Value'].sum() + future_df['Present Value'].sum()):,.2f}
• Adjusted Earnings Factor: {aef_result['final_aef']:.3f}
"""
    
    doc.add_paragraph(summary_text)
    
    # Case Information
    doc.add_heading('CASE INFORMATION', level=1)
    
    # Create case info table
    case_table = doc.add_table(rows=12, cols=2)
    case_table.style = 'Table Grid'
    
    case_info_data = [
        ('Client Name', data.client_name),
        ('Date of Birth', data.date_of_birth.strftime('%m/%d/%Y')),
        ('Date of Injury', data.date_of_injury.strftime('%m/%d/%Y')),
        ('Age at Injury', f"{data.age_at_injury:.1f} years"),
        ('Gender', data.gender),
        ('Education', data.education),
        ('Pre-Injury Occupation', data.pre_injury_occupation),
        ('Pre-Injury Annual Income', f"${data.pre_injury_annual_income:,.2f}"),
        ('Life Expectancy', f"{data.life_expectancy:.1f} years"),
        ('Work Life Expectancy', f"{data.work_life_expectancy:.1f} years"),
        ('Statistical Retirement Date', retirement_date.strftime('%m/%d/%Y')),
        ('Statistical Death Date', death_date.strftime('%m/%d/%Y'))
    ]
    
    for i, (label, value) in enumerate(case_info_data):
        case_table.cell(i, 0).text = label
        case_table.cell(i, 1).text = str(value)
    
    # Economic Assumptions
    doc.add_heading('ECONOMIC ASSUMPTIONS', level=1)
    
    assumptions_table = doc.add_table(rows=6, cols=2)
    assumptions_table.style = 'Table Grid'
    
    assumptions_data = [
        ('Wage Growth Rate', f"{data.wage_growth_rate * 100:.1f}%"),
        ('Discount Rate', f"{data.discount_rate * 100:.1f}%"),
        ('Unemployment Rate', f"{data.unemployment_rate * 100:.1f}%"),
        ('Tax Rate', f"{data.tax_rate * 100:.1f}%"),
        ('Fringe Benefits Rate', f"{data.fringe_benefits_rate * 100:.1f}%"),
        ('Personal Consumption Rate', f"{data.personal_consumption_rate * 100:.1f}%")
    ]
    
    for i, (label, value) in enumerate(assumptions_data):
        assumptions_table.cell(i, 0).text = label
        assumptions_table.cell(i, 1).text = value
    
    # AEF Breakdown
    doc.add_heading('ADJUSTED EARNINGS FACTOR (AEF)', level=1)
    
    aef_table = doc.add_table(rows=len(aef_result['factors']) + 1, cols=3)
    aef_table.style = 'Table Grid'
    
    # Headers
    aef_table.cell(0, 0).text = 'Factor'
    aef_table.cell(0, 1).text = 'Value'
    aef_table.cell(0, 2).text = 'Cumulative'
    
    for i, (factor_name, factor_value, cumulative) in enumerate(aef_result['factors']):
        aef_table.cell(i + 1, 0).text = factor_name
        aef_table.cell(i + 1, 1).text = f"{factor_value:.3f}"
        aef_table.cell(i + 1, 2).text = f"{cumulative:.3f}"
    
    # Past Losses Table
    doc.add_page_break()
    doc.add_heading('PAST ECONOMIC LOSSES', level=1)
    
    if not past_df.empty:
        past_table = doc.add_table(rows=len(past_df) + 2, cols=len(past_df.columns))
        past_table.style = 'Table Grid'
        
        # Headers
        for j, col in enumerate(past_df.columns):
            past_table.cell(0, j).text = col
        
        # Data
        for i in range(len(past_df)):
            for j, col in enumerate(past_df.columns):
                value = past_df.iloc[i, j]
                if isinstance(value, (int, float)) and col in ['Pre-Injury Earnings', 'Nominal Loss', 'Adjusted Loss', 'Present Value']:
                    past_table.cell(i + 1, j).text = f"${value:,.2f}"
                else:
                    past_table.cell(i + 1, j).text = str(value)
        
        # Totals row
        totals_row = len(past_df) + 1
        past_table.cell(totals_row, 0).text = "TOTALS"
        for j, col in enumerate(past_df.columns):
            if col in ['Pre-Injury Earnings', 'Nominal Loss', 'Adjusted Loss', 'Present Value']:
                total = past_df[col].sum()
                past_table.cell(totals_row, j).text = f"${total:,.2f}"
    
    # Future Losses Table
    doc.add_page_break()
    doc.add_heading('FUTURE ECONOMIC LOSSES', level=1)
    
    if not future_df.empty:
        future_table = doc.add_table(rows=len(future_df) + 2, cols=len(future_df.columns))
        future_table.style = 'Table Grid'
        
        # Headers
        for j, col in enumerate(future_df.columns):
            future_table.cell(0, j).text = col
        
        # Data
        for i in range(len(future_df)):
            for j, col in enumerate(future_df.columns):
                value = future_df.iloc[i, j]
                if isinstance(value, (int, float)) and col in ['Pre-Injury Earnings', 'Post-Injury Earnings', 'Nominal Loss', 'Adjusted Loss', 'Present Value']:
                    future_table.cell(i + 1, j).text = f"${value:,.2f}"
                else:
                    future_table.cell(i + 1, j).text = str(value)
        
        # Totals row
        totals_row = len(future_df) + 1
        future_table.cell(totals_row, 0).text = "TOTALS"
        for j, col in enumerate(future_df.columns):
            if col in ['Pre-Injury Earnings', 'Post-Injury Earnings', 'Nominal Loss', 'Adjusted Loss', 'Present Value']:
                total = future_df[col].sum()
                future_table.cell(totals_row, j).text = f"${total:,.2f}"
    
    # Summary
    doc.add_page_break()
    doc.add_heading('SUMMARY OF ECONOMIC LOSSES', level=1)
    
    summary_table = doc.add_table(rows=4, cols=2)
    summary_table.style = 'Table Grid'
    
    past_total = past_df['Present Value'].sum() if not past_df.empty else 0
    future_total = future_df['Present Value'].sum() if not future_df.empty else 0
    total_loss = past_total + future_total
    
    summary_data = [
        ('Past Economic Losses', f"${past_total:,.2f}"),
        ('Future Economic Losses', f"${future_total:,.2f}"),
        ('Life Care Plan Costs', f"${data.life_care_plan_cost:,.2f}"),
        ('TOTAL ECONOMIC LOSS', f"${total_loss + data.life_care_plan_cost:,.2f}")
    ]
    
    for i, (label, value) in enumerate(summary_data):
        summary_table.cell(i, 0).text = label
        summary_table.cell(i, 1).text = value
        if 'TOTAL' in label:
            # Make total row bold
            for paragraph in summary_table.cell(i, 0).paragraphs:
                for run in paragraph.runs:
                    run.bold = True
            for paragraph in summary_table.cell(i, 1).paragraphs:
                for run in paragraph.runs:
                    run.bold = True
    
    # Save to bytes
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io.read()

# Streamlit Interface
def main():
    # Sidebar for main inputs
    with st.sidebar:
        st.header("📋 Report Configuration")
        
        # Load/Save functionality
        st.subheader("Data Management")
        
        col1, col2 = st.columns(2)
        with col1:
            if st.button("💾 Save Data"):
                data_json = json.dumps(st.session_state.report_data, default=str)
                st.download_button(
                    "Download JSON",
                    data_json,
                    f"economic_report_data_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                    "application/json"
                )
        
        with col2:
            uploaded_file = st.file_uploader("📁 Load Data", type=['json'])
            if uploaded_file:
                try:
                    loaded_data = json.load(uploaded_file)
                    st.session_state.report_data.update(loaded_data)
                    st.success("Data loaded successfully!")
                except Exception as e:
                    st.error(f"Error loading file: {str(e)}")
    
    # Main content area with tabs
    tab1, tab2, tab3, tab4, tab5 = st.tabs([
        "👤 Personal Info", 
        "💼 Employment", 
        "🏥 Injury Details", 
        "📊 Economic Factors", 
        "📋 Generate Report"
    ])
    
    # Initialize data object
    data = EconomicData()
    
    with tab1:
        st.header("Personal Information")
        
        col1, col2 = st.columns(2)
        with col1:
            data.client_name = st.text_input(
                "Client Name", 
                value=st.session_state.report_data.get('client_name', ''),
                key='client_name'
            )
            data.date_of_birth = st.date_input(
                "Date of Birth", 
                value=st.session_state.report_data.get('date_of_birth', date(1980, 1, 1)),
                key='date_of_birth'
            )
            data.gender = st.selectbox(
                "Gender", 
                ["Male", "Female"],
                index=0 if st.session_state.report_data.get('gender', 'Male') == 'Male' else 1,
                key='gender'
            )
            data.education = st.text_input(
                "Education Level", 
                value=st.session_state.report_data.get('education', ''),
                key='education'
            )
        
        with col2:
            data.date_of_injury = st.date_input(
                "Date of Injury", 
                value=st.session_state.report_data.get('date_of_injury', date(2023, 1, 1)),
                key='date_of_injury'
            )
            data.date_of_report = st.date_input(
                "Date of Report", 
                value=st.session_state.report_data.get('date_of_report', date.today()),
                key='date_of_report'
            )
            data.marital_status = st.text_input(
                "Marital Status", 
                value=st.session_state.report_data.get('marital_status', ''),
                key='marital_status'
            )
        
        # Calculate age at injury
        data.age_at_injury = calculate_age(data.date_of_birth, data.date_of_injury)
        st.info(f"Age at injury: {data.age_at_injury:.1f} years")
        
        st.subheader("Geographic Information")
        col3, col4, col5 = st.columns(3)
        with col3:
            data.residence_state = st.text_input(
                "State of Residence", 
                value=st.session_state.report_data.get('residence_state', ''),
                key='residence_state'
            )
        with col4:
            data.residence_county = st.text_input(
                "County", 
                value=st.session_state.report_data.get('residence_county', ''),
                key='residence_county'
            )
        with col5:
            data.msa = st.text_input(
                "MSA", 
                value=st.session_state.report_data.get('msa', ''),
                key='msa'
            )
    
    with tab2:
        st.header("Employment Information")
        
        col1, col2 = st.columns(2)
        with col1:
            data.pre_injury_occupation = st.text_input(
                "Pre-Injury Occupation", 
                value=st.session_state.report_data.get('pre_injury_occupation', ''),
                key='pre_injury_occupation'
            )
            data.pre_injury_annual_income = st.number_input(
                "Pre-Injury Annual Income ($)", 
                min_value=0.0, 
                value=float(st.session_state.report_data.get('pre_injury_annual_income', 50000.0)),
                step=1000.0,
                key='pre_injury_annual_income'
            )
            data.employer = st.text_input(
                "Employer", 
                value=st.session_state.report_data.get('employer', ''),
                key='employer'
            )
        
        with col2:
            data.pre_injury_hourly_rate = st.number_input(
                "Pre-Injury Hourly Rate ($)", 
                min_value=0.0, 
                value=float(st.session_state.report_data.get('pre_injury_hourly_rate', 24.04)),
                step=0.50,
                key='pre_injury_hourly_rate'
            )
            data.years_with_employer = st.number_input(
                "Years with Employer", 
                min_value=0.0, 
                value=float(st.session_state.report_data.get('years_with_employer', 5.0)),
                step=0.5,
                key='years_with_employer'
            )
        
        st.subheader("Post-Injury Employment")
        col3, col4 = st.columns(2)
        with col3:
            data.post_injury_occupation = st.text_input(
                "Post-Injury Occupation", 
                value=st.session_state.report_data.get('post_injury_occupation', ''),
                key='post_injury_occupation'
            )
            data.post_injury_capacity = st.text_area(
                "Post-Injury Work Capacity", 
                value=st.session_state.report_data.get('post_injury_capacity', ''),
                key='post_injury_capacity'
            )
        
        with col4:
            data.post_injury_annual_income = st.number_input(
                "Post-Injury Annual Income ($)", 
                min_value=0.0, 
                value=float(st.session_state.report_data.get('post_injury_annual_income', 0.0)),
                step=1000.0,
                key='post_injury_annual_income'
            )
            data.labor_market_reduction = st.number_input(
                "Labor Market Reduction (%)", 
                min_value=0.0,
                max_value=100.0,
                value=float(st.session_state.report_data.get('labor_market_reduction', 0.0)),
                step=1.0,
                key='labor_market_reduction'
            ) / 100
    
    with tab3:
        st.header("Injury and Medical Information")
        
        col1, col2 = st.columns(2)
        with col1:
            data.injury_description = st.text_area(
                "Injury Description", 
                value=st.session_state.report_data.get('injury_description', ''),
                key='injury_description',
                height=100
            )
            data.body_parts_injured = st.text_input(
                "Body Parts Injured", 
                value=st.session_state.report_data.get('body_parts_injured', ''),
                key='body_parts_injured'
            )
        
        with col2:
            data.medical_treatment = st.text_area(
                "Medical Treatment", 
                value=st.session_state.report_data.get('medical_treatment', ''),
                key='medical_treatment',
                height=100
            )
            data.current_medical_status = st.text_area(
                "Current Medical Status", 
                value=st.session_state.report_data.get('current_medical_status', ''),
                key='current_medical_status',
                height=100
            )
    
    with tab4:
        st.header("Economic Factors and Life Expectancy")
        
        st.subheader("Life Expectancy Data")
        col1, col2, col3 = st.columns(3)
        with col1:
            data.life_expectancy = st.number_input(
                "Life Expectancy (years)", 
                min_value=0.0,
                max_value=120.0,
                value=float(st.session_state.report_data.get('life_expectancy', 78.5)),
                step=0.1,
                key='life_expectancy'
            )
        with col2:
            data.work_life_expectancy = st.number_input(
                "Work Life Expectancy (years)", 
                min_value=0.0,
                max_value=80.0,
                value=float(st.session_state.report_data.get('work_life_expectancy', 45.0)),
                step=0.1,
                key='work_life_expectancy'
            )
        with col3:
            data.years_to_final_separation = st.number_input(
                "Years to Final Separation", 
                min_value=0.0,
                value=float(st.session_state.report_data.get('years_to_final_separation', 0.0)),
                step=0.1,
                key='years_to_final_separation'
            )
        
        # Calculate and display statistical dates
        retirement_date, death_date = calculate_statistical_dates(
            data.date_of_injury, data.life_expectancy, data.work_life_expectancy
        )
        
        col4, col5 = st.columns(2)
        with col4:
            st.info(f"Statistical Retirement Date: {retirement_date.strftime('%m/%d/%Y')}")
        with col5:
            st.info(f"Statistical Death Date: {death_date.strftime('%m/%d/%Y')}")
        
        st.subheader("Economic Growth and Discount Rates")
        col6, col7 = st.columns(2)
        with col6:
            data.wage_growth_rate = st.number_input(
                "Wage Growth Rate (%)", 
                min_value=0.0,
                max_value=20.0,
                value=float(st.session_state.report_data.get('wage_growth_rate', 3.0)),
                step=0.1,
                key='wage_growth_rate'
            ) / 100
        with col7:
            data.discount_rate = st.number_input(
                "Discount Rate (%)", 
                min_value=0.0,
                max_value=20.0,
                value=float(st.session_state.report_data.get('discount_rate', 4.0)),
                step=0.1,
                key='discount_rate'
            ) / 100
        
        st.subheader("Adjustment Factors")
        col8, col9 = st.columns(2)
        with col8:
            data.unemployment_rate = st.number_input(
                "Unemployment Rate (%)", 
                min_value=0.0,
                max_value=30.0,
                value=float(st.session_state.report_data.get('unemployment_rate', 3.5)),
                step=0.1,
                key='unemployment_rate'
            ) / 100
            data.tax_rate = st.number_input(
                "Tax Rate (%)", 
                min_value=0.0,
                max_value=50.0,
                value=float(st.session_state.report_data.get('tax_rate', 12.0)),
                step=0.1,
                key='tax_rate'
            ) / 100
        with col9:
            data.fringe_benefits_rate = st.number_input(
                "Fringe Benefits Rate (%)", 
                min_value=0.0,
                max_value=50.0,
                value=float(st.session_state.report_data.get('fringe_benefits_rate', 6.0)),
                step=0.1,
                key='fringe_benefits_rate'
            ) / 100
            data.personal_consumption_rate = st.number_input(
                "Personal Consumption Rate (%)", 
                min_value=0.0,
                max_value=50.0,
                value=float(st.session_state.report_data.get('personal_consumption_rate', 0.0)),
                step=0.1,
                key='personal_consumption_rate',
                help="Use for wrongful death cases"
            ) / 100
        
        st.subheader("Additional Costs")
        col10, col11 = st.columns(2)
        with col10:
            data.life_care_plan_cost = st.number_input(
                "Life Care Plan Cost ($)", 
                min_value=0.0,
                value=float(st.session_state.report_data.get('life_care_plan_cost', 0.0)),
                step=1000.0,
                key='life_care_plan_cost'
            )
        with col11:
            data.household_services_cost = st.number_input(
                "Household Services Cost ($)", 
                min_value=0.0,
                value=float(st.session_state.report_data.get('household_services_cost', 0.0)),
                step=1000.0,
                key='household_services_cost'
            )
    
    with tab5:
        st.header("Generate Economic Loss Report")
        
        # Update session state with current data
        for key, value in data.__dict__.items():
            st.session_state.report_data[key] = value
        
        if st.button("🔄 Calculate Economic Losses", type="primary"):
            try:
                # Generate calculations
                past_df, future_df, aef_result, retirement_date, death_date = generate_economic_schedule(data)
                
                # Display results
                st.subheader("📊 Calculation Results")
                
                # Summary metrics
                past_total = past_df['Present Value'].sum() if not past_df.empty else 0
                future_total = future_df['Present Value'].sum() if not future_df.empty else 0
                total_loss = past_total + future_total + data.life_care_plan_cost
                
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    st.metric("Past Losses", f"${past_total:,.2f}")
                with col2:
                    st.metric("Future Losses", f"${future_total:,.2f}")
                with col3:
                    st.metric("Life Care Plan", f"${data.life_care_plan_cost:,.2f}")
                with col4:
                    st.metric("Total Economic Loss", f"${total_loss:,.2f}")
                
                # AEF breakdown
                st.subheader("🧮 Adjusted Earnings Factor (AEF)")
                aef_data = []
                for factor_name, factor_value, cumulative in aef_result['factors']:
                    aef_data.append({
                        'Factor': factor_name,
                        'Value': f"{factor_value:.3f}",
                        'Percentage': f"{factor_value * 100:.1f}%",
                        'Cumulative': f"{cumulative:.3f}"
                    })
                
                aef_df = pd.DataFrame(aef_data)
                st.dataframe(aef_df, use_container_width=True, hide_index=True)
                
                # Past losses table
                if not past_df.empty:
                    st.subheader("📅 Past Economic Losses")
                    st.dataframe(past_df, use_container_width=True, hide_index=True)
                
                # Future losses table
                if not future_df.empty:
                    st.subheader("🔮 Future Economic Losses")
                    st.dataframe(future_df, use_container_width=True, hide_index=True)
                
                # Generate Word report
                st.subheader("📄 Download Report")
                
                if st.button("Generate Word Report"):
                    with st.spinner("Generating Word document..."):
                        doc_bytes = create_word_report(data, past_df, future_df, aef_result, retirement_date, death_date)
                        
                        st.download_button(
                            label="📄 Download Word Report",
                            data=doc_bytes,
                            file_name=f"Economic_Loss_Report_{data.client_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                        
                        st.success("Word report generated successfully!")
                
                # Export data options
                st.subheader("📊 Export Data")
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    if not past_df.empty:
                        csv_past = past_df.to_csv(index=False)
                        st.download_button(
                            "📊 Past Losses CSV",
                            csv_past,
                            f"past_losses_{data.client_name.replace(' ', '_')}.csv",
                            "text/csv"
                        )
                
                with col2:
                    if not future_df.empty:
                        csv_future = future_df.to_csv(index=False)
                        st.download_button(
                            "📊 Future Losses CSV",
                            csv_future,
                            f"future_losses_{data.client_name.replace(' ', '_')}.csv",
                            "text/csv"
                        )
                
                with col3:
                    # Combined Excel export
                    excel_io = io.BytesIO()
                    with pd.ExcelWriter(excel_io, engine='openpyxl') as writer:
                        if not past_df.empty:
                            past_df.to_excel(writer, sheet_name='Past Losses', index=False)
                        if not future_df.empty:
                            future_df.to_excel(writer, sheet_name='Future Losses', index=False)
                        aef_df.to_excel(writer, sheet_name='AEF Breakdown', index=False)
                    
                    excel_io.seek(0)
                    st.download_button(
                        "📊 Complete Excel Report",
                        excel_io,
                        f"economic_analysis_{data.client_name.replace(' ', '_')}.xlsx",
                        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    )
                
            except Exception as e:
                st.error(f"Error in calculations: {str(e)}")
                st.exception(e)

if __name__ == "__main__":
    main()